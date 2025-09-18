
import os
import sys
import io
import json
import queue
import threading
import tempfile
import subprocess
import base64
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple, Union

# --- OCR & PDF ---
import fitz  # PyMuPDF
from PIL import Image
try:
    from pdf2image import convert_from_path  # Para convers√£o de PDF em imagens
    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False

# --- HTTP & env ---
import requests
from dotenv import load_dotenv
from datetime import datetime

# --- Utilidades ---
import textwrap

# --- Exporta√ß√£o de documentos ---
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# --- GUI ---
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinter.font as tkfont

# --- Auto-atualiza√ß√£o ---
from updater import create_updater
from feedback_system import initialize_feedback_system, get_feedback_system

# =========================u
# Configura√ß√£o
# =========================
APP_TITLE = "Analisador de Usucapi√£o com IA Visual ‚Äì Matr√≠culas e Confrontantes (PGE-MS)"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

# Carrega .env
load_dotenv()


def _load_app_version(default="1.0.0"):
    """Recupera versao do arquivo VERSION ou retorna padrao"""
    version_file = Path(__file__).parent.parent / "VERSION"
    try:
        content = version_file.read_text(encoding="utf-8").strip()
        return content or default
    except (OSError, UnicodeDecodeError):
        return default


APP_VERSION = _load_app_version()

DEFAULT_MODEL = os.environ.get("OPENROUTER_MODEL", "google/gemini-2.5-pro")
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
FULL_REPORT_MODEL = "google/gemini-2.5-flash"

# Configura√ß√£o do Google Forms para Feedback
GOOGLE_FORM_CONFIG = {
    "url": os.getenv("GOOGLE_FORM_URL", ""),
    "fields": {
        "tipo": os.getenv("GOOGLE_FORM_FIELD_TIPO", ""),
        "descricao": os.getenv("GOOGLE_FORM_FIELD_DESCRICAO", ""),
        "modelo": os.getenv("GOOGLE_FORM_FIELD_MODELO", ""),
        "timestamp": os.getenv("GOOGLE_FORM_FIELD_TIMESTAMP", ""),
        "versao": os.getenv("GOOGLE_FORM_FIELD_VERSAO", "")
    }
}

# =========================
# Estruturas
# =========================
@dataclass
class TransmissaoInfo:
    """Informa√ß√µes sobre uma transmiss√£o na cadeia dominial"""
    data: Optional[str] = None
    tipo_transmissao: Optional[str] = None
    proprietario_anterior: Optional[str] = None
    novo_proprietario: Optional[str] = None
    percentual: Optional[str] = None
    valor: Optional[str] = None
    registro: Optional[str] = None

@dataclass
class RestricaoInfo:
    """Informa√ß√µes sobre restri√ß√µes e gravames"""
    tipo: str
    data_registro: Optional[str] = None
    credor: Optional[str] = None
    valor: Optional[str] = None
    situacao: str = "vigente"  # "vigente" ou "baixada"
    data_baixa: Optional[str] = None
    observacoes: Optional[str] = None


@dataclass
class MatriculaInfo:
    numero: str
    proprietarios: List[str]
    descricao: str
    confrontantes: List[str]
    evidence: List[str]
    lote: Optional[str] = None  # n√∫mero do lote
    quadra: Optional[str] = None  # n√∫mero da quadra
    cadeia_dominial: List[TransmissaoInfo] = None  # hist√≥rico de transmiss√µes
    restricoes: List[RestricaoInfo] = None  # restri√ß√µes e gravames
    
    def __post_init__(self):
        if self.cadeia_dominial is None:
            self.cadeia_dominial = []
        if self.restricoes is None:
            self.restricoes = []

@dataclass
class LoteConfronta:
    """Informa√ß√µes sobre um lote confrontante"""
    identificador: str  # "lote 10", "matr√≠cula 1234", etc.
    tipo: str  # "lote", "matr√≠cula", "pessoa", "via_publica", "estado", "outros"
    matricula_anexada: Optional[str] = None  # n√∫mero da matr√≠cula se foi anexada
    direcao: Optional[str] = None  # norte, sul, leste, oeste, etc.
    
@dataclass
class EstadoMSDireitos:
    """Informa√ß√µes sobre direitos do Estado de MS"""
    tem_direitos: bool = False
    detalhes: List[Dict] = None
    criticidade: str = "baixa"  # "alta", "media", "baixa"
    observacao: str = ""
    
    def __post_init__(self):
        if self.detalhes is None:
            self.detalhes = []

@dataclass
class ResumoAnalise:
    """Resumo estruturado da an√°lise para o relat√≥rio"""
    cadeia_dominial_completa: Dict[str, List[Dict]] = None  # matr√≠cula -> lista cronol√≥gica
    restricoes_vigentes: List[Dict] = None  # restri√ß√µes ainda em vigor
    restricoes_baixadas: List[Dict] = None  # restri√ß√µes j√° canceladas
    estado_ms_direitos: EstadoMSDireitos = None  # direitos do Estado de MS
    
    def __post_init__(self):
        if self.cadeia_dominial_completa is None:
            self.cadeia_dominial_completa = {}
        if self.restricoes_vigentes is None:
            self.restricoes_vigentes = []
        if self.restricoes_baixadas is None:
            self.restricoes_baixadas = []
        if self.estado_ms_direitos is None:
            self.estado_ms_direitos = EstadoMSDireitos()

@dataclass
class AnalysisResult:
    arquivo: str
    matriculas_encontradas: List[MatriculaInfo]
    matricula_principal: Optional[str]  # n√∫mero da matr√≠cula de usucapi√£o
    matriculas_confrontantes: List[str]  # n√∫meros das matr√≠culas confrontantes
    # NOVOS CAMPOS PARA MELHOR CONTROLE
    lotes_confrontantes: List[LoteConfronta]  # todos os confrontantes identificados
    matriculas_nao_confrontantes: List[str]  # matr√≠culas anexadas que N√ÉO s√£o confrontantes
    lotes_sem_matricula: List[str]  # lotes confrontantes sem matr√≠cula anexada
    confrontacao_completa: Optional[bool]  # se todas confrontantes foram apresentadas
    proprietarios_identificados: Dict[str, List[str]]  # n√∫mero -> lista propriet√°rios
    resumo_analise: Optional[ResumoAnalise] = None  # resumo estruturado da an√°lise
    confidence: Optional[float] = None
    reasoning: str = ""
    raw_json: Dict = None
    
    def __post_init__(self):
        if self.resumo_analise is None:
            self.resumo_analise = ResumoAnalise()
        if self.raw_json is None:
            self.raw_json = {}
    
    # Campos de compatibilidade (para n√£o quebrar c√≥digo existente)
    @property
    def is_confrontante(self) -> Optional[bool]:
        """Compatibilidade: retorna se encontrou Estado MS como confrontante"""
        estado_patterns = ['estado de mato grosso do sul', 'estado de ms', 'estado do ms', 
                          'fazenda do estado', 'governo do estado', 'fazenda p√∫blica estadual']
        for matricula in self.matriculas_encontradas:
            for confrontante in matricula.confrontantes:
                if any(pattern in confrontante.lower() for pattern in estado_patterns):
                    return True
        return False
    


def image_to_base64(image_path_or_pil: Union[str, Image.Image], max_size: int = 1024, jpeg_quality: int = 85) -> str:
    """
    Converte imagem para base64 otimizada para envio √† API de vis√£o.
    """
    try:
        if isinstance(image_path_or_pil, str):
            img = Image.open(image_path_or_pil)
        else:
            img = image_path_or_pil
        
        # Converte para RGB se necess√°rio
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Redimensiona se muito grande (mant√©m propor√ß√£o)
        if max(img.size) > max_size:
            img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        
        # Converte para base64
        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', quality=jpeg_quality, optimize=True)
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return img_str
    except Exception as e:
        print(f"Erro ao converter imagem para base64: {e}")
        return ""

def get_pdf_page_count(pdf_path: str) -> int:
    """
    Retorna o n√∫mero total de p√°ginas de um PDF.
    """
    try:
        if PDF2IMAGE_AVAILABLE:
            try:
                from pdf2image.utils import get_page_count  # type: ignore
                return get_page_count(pdf_path)
            except Exception:
                pass
        
        # Fallback com PyMuPDF
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        doc.close()
        return page_count
        
    except Exception as e:
        print(f"Erro ao contar p√°ginas do PDF: {e}")
        return 0

def pdf_to_images(pdf_path: str, max_pages: Optional[int] = 10) -> List[Image.Image]:
    """
    Converte PDF para lista de imagens PIL para an√°lise visual.
    Se max_pages for None, processa todas as p√°ginas.
    """
    images = []
    try:
        # Primeiro tenta com pdf2image (mais r√°pido)
        if PDF2IMAGE_AVAILABLE:
            try:
                if max_pages is None:
                    # Sem limite - processa todas as p√°ginas
                    pdf_images = convert_from_path(pdf_path, dpi=200)
                else:
                    pdf_images = convert_from_path(pdf_path, dpi=200, first_page=1, last_page=max_pages)
                return pdf_images
            except Exception:
                pass
        
        # Fallback com PyMuPDF
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        pages_to_process = total_pages if max_pages is None else min(total_pages, max_pages)
        
        for page_num in range(pages_to_process):
            page = doc[page_num]
            # Converte p√°gina para imagem
            mat = fitz.Matrix(2.0, 2.0)  # escala 2x para melhor qualidade
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            images.append(img)
        doc.close()
        
    except Exception as e:
        print(f"Erro ao converter PDF para imagens: {e}")
    
    return images


# =========================
# Cliente OpenRouter
# =========================
def call_openrouter_vision(model: str, system_prompt: str, user_prompt: str, images_base64: List[str], temperature: float = 0.0, max_tokens: int = 1500) -> Dict:
    """
    Chama a API OpenRouter com suporte a vis√£o computacional (an√°lise de imagens).
    """
    api_key = os.environ.get("OPENROUTER_API_KEY", OPENROUTER_API_KEY)
    if not api_key:
        raise RuntimeError("OPENROUTER_API_KEY n√£o configurada.")

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://pge-ms.lab/analise-matriculas",
        "X-Title": "Analise de Matriculas PGE-MS"
    }

    # Constr√≥i mensagem com imagens
    content = [{"type": "text", "text": user_prompt}]
    
    # Adiciona cada imagem
    for i, img_b64 in enumerate(images_base64):
        if img_b64:  # verifica se base64 n√£o est√° vazio
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{img_b64}",
                    "detail": "high"  # alta qualidade para documentos
                }
            })

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ],
        "temperature": 0.1,  # Reduzido para respostas mais focadas
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object"}
    }

    try:
        # Debug detalhado do payload
        try:
            message_content = payload['messages'][1]['content']
            image_count = sum(1 for item in message_content if item.get('type') == 'image_url')
            text_count = sum(1 for item in message_content if item.get('type') == 'text')
            
            print(f"üåê Fazendo requisi√ß√£o para: {OPENROUTER_URL}")
            print(f"üì¶ Payload cont√©m {len(message_content)} elementos total")
            print(f"üñºÔ∏è Imagens no payload: {image_count}")
            print(f"üìù Textos no payload: {text_count}")
            print(f"üîë Modelo: {payload.get('model', 'N/A')}")
            
            # Calcula tamanho total do payload em MB
            import sys
            payload_size_mb = sys.getsizeof(str(payload)) / (1024 * 1024)
            print(f"üìê Tamanho do payload: {payload_size_mb:.2f}MB")
            
        except Exception as e:
            print(f"‚ö†Ô∏è Erro ao analisar payload: {e}")
            print(f"üìä Estrutura do payload: {list(payload.keys()) if isinstance(payload, dict) else type(payload)}")
        
        resp = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=120)
        
        print(f"üì° Status da resposta: {resp.status_code}")
        print(f"üìä Headers da resposta: {dict(list(resp.headers.items())[:5])}...")  # primeiros 5 headers
        
        if resp.status_code != 200:
            print(f"‚ùå Erro HTTP {resp.status_code}: {resp.text[:500]}")
            # Tenta extrair mais detalhes do erro
            try:
                error_data = json.loads(resp.text)
                if "error" in error_data:
                    error_msg = error_data["error"]
                    if isinstance(error_msg, dict):
                        error_details = error_msg.get("message", str(error_msg))
                    else:
                        error_details = str(error_msg)
                    raise RuntimeError(f"API Error ({resp.status_code}): {error_details}")
            except json.JSONDecodeError:
                pass
            raise RuntimeError(f"API retornou status {resp.status_code}: {resp.text[:200]}")
            
        response_text = resp.text.strip()
        print(f"üìù Tamanho da resposta: {len(response_text)} chars")
        
        if not response_text:
            raise RuntimeError("Resposta vazia da API")
        
        # Debug da resposta bruta
        if len(response_text) < 200:
            print(f"üìÑ Resposta completa: {response_text}")
        else:
            print(f"üìÑ In√≠cio da resposta: {response_text[:300]}...")
            print(f"üìÑ Final da resposta: ...{response_text[-100:]}")
            
        # Parse mais robusto do JSON
        try:
            data = json.loads(response_text)
        except json.JSONDecodeError as e:
            print(f"‚ùå Erro JSON: {e}")
            print(f"üìÑ Conte√∫do problem√°tico: {response_text[:1000]}")
            raise RuntimeError(f"Resposta da API n√£o √© JSON v√°lido: {e}")
        
        # Debug da estrutura da resposta
        print(f"üîç Estrutura da resposta: {list(data.keys()) if isinstance(data, dict) else type(data)}")
        
        if not isinstance(data, dict):
            raise RuntimeError(f"Resposta da API n√£o √© um objeto JSON: {type(data)}")
        
        if "choices" not in data:
            print(f"‚ùå Campo 'choices' n√£o encontrado. Campos dispon√≠veis: {list(data.keys())}")
            # Verifica se h√° uma mensagem de erro
            if "error" in data:
                error_msg = data["error"]
                raise RuntimeError(f"API retornou erro: {error_msg}")
            raise RuntimeError(f"Campo 'choices' ausente na resposta. Estrutura: {data}")
        
        if not data["choices"]:
            print(f"‚ùå Lista 'choices' est√° vazia")
            raise RuntimeError("Lista 'choices' vazia na resposta da API")
        
        if not isinstance(data["choices"], list):
            print(f"‚ùå 'choices' n√£o √© uma lista: {type(data['choices'])}")
            raise RuntimeError(f"Campo 'choices' deve ser uma lista, mas √©: {type(data['choices'])}")
        
        print(f"‚úÖ Resposta v√°lida com {len(data['choices'])} choice(s)")
        return data
        
    except requests.exceptions.RequestException as e:
        raise RuntimeError(f"Erro na requisi√ß√£o para OpenRouter: {e}")
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Erro ao decodificar JSON da resposta: {e}. Resposta: {response_text[:500]}")
    except Exception as e:
        raise RuntimeError(f"Erro inesperado na chamada da API: {e}")


def call_openrouter_text(model: str, system_prompt: str, user_prompt: str, temperature: float = 0.2, max_tokens: int = 2000) -> str:
    """Chama a API OpenRouter para gerar texto com base em prompt estruturado."""
    api_key = os.environ.get("OPENROUTER_API_KEY", OPENROUTER_API_KEY)
    if not api_key:
        raise RuntimeError("OPENROUTER_API_KEY n√£o configurada.")

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://pge-ms.lab/analise-matriculas",
        "X-Title": "Analise de Matriculas PGE-MS"
    }

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": temperature,
        "max_tokens": max_tokens
    }

    try:
        print(f"üåê [Texto] Requisi√ß√£o para {OPENROUTER_URL} com modelo {model}")
        resp = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=120)
        print(f"üì° [Texto] Status: {resp.status_code}")

        if resp.status_code != 200:
            preview = resp.text[:500]
            raise RuntimeError(f"API retornou status {resp.status_code}: {preview}")

        data = resp.json()
        if not isinstance(data, dict) or "choices" not in data or not data["choices"]:
            raise RuntimeError(f"Resposta inesperada da API: {data}")

        message = data["choices"][0]["message"].get("content", "")
        if not message:
            raise RuntimeError("Resposta da API n√£o cont√©m conte√∫do textual.")

        print(f"‚úÖ [Texto] Conte√∫do recebido com {len(message)} caracteres")
        return message

    except requests.exceptions.RequestException as exc:
        raise RuntimeError(f"Erro na requisi√ß√£o para OpenRouter: {exc}") from exc


def clean_json_response(content: str) -> str:
    """Extrai JSON de uma resposta que pode conter markdown e texto adicional"""
    content = content.strip()
    
    # Procura por blocos JSON em markdown
    import re
    
    # Padr√£o 1: ```json ... ```
    json_pattern = r'```json\s*\n(.*?)\n```'
    match = re.search(json_pattern, content, re.DOTALL)
    if match:
        json_content = match.group(1).strip()
        print(f"‚úÖ JSON extra√≠do do markdown (```json): {len(json_content)} chars")
        return json_content
    
    # Padr√£o 2: ``` ... ``` (sem especificar json)
    json_pattern = r'```\s*\n(.*?)\n```'
    match = re.search(json_pattern, content, re.DOTALL)
    if match:
        candidate = match.group(1).strip()
        # Verifica se parece com JSON (come√ßa com { ou [)
        if candidate.startswith('{') or candidate.startswith('['):
            print(f"‚úÖ JSON extra√≠do do markdown (```): {len(candidate)} chars")
            return candidate
    
    # Padr√£o 3: Procura por { ... } que parece ser JSON
    json_pattern = r'\{.*\}'
    match = re.search(json_pattern, content, re.DOTALL)
    if match:
        candidate = match.group(0).strip()
        print(f"‚úÖ JSON extra√≠do por regex {{...}}: {len(candidate)} chars")
        return candidate
    
    # Se n√£o encontrou nada, retorna o conte√∫do original
    print(f"‚ö†Ô∏è Nenhum JSON encontrado, retornando conte√∫do original: {len(content)} chars")
    return content

# =========================
# Prompting
# =========================

# Sistema unificado de prompts para an√°lise de matr√≠culas imobili√°rias
UNIFIED_SYSTEM_PROMPT = (
    "Voc√™ √© um perito ESPECIALISTA em an√°lise de processos de usucapi√£o e matr√≠culas imobili√°rias brasileiras. "
    "Sua responsabilidade √© CR√çTICA: a identifica√ß√£o COMPLETA de confrontantes pode determinar o sucesso ou fracasso de um usucapi√£o.\n\n"

    "üéØ MISS√ÉO VITAL:\n"
    "‚Ä¢ IDENTIFIQUE TODOS os confrontantes da matr√≠cula principal SEM EXCE√á√ÉO\n"
    "‚Ä¢ TODO LOTE DEVE TER NO M√çNIMO 4 CONFRONTANTES (uma para cada dire√ß√£o)\n"
    "‚Ä¢ EXTRAIA LITERALMENTE cada nome, matr√≠cula, rua mencionada como confrontante\n"
    "‚Ä¢ ANALISE palavra por palavra a descri√ß√£o do im√≥vel principal\n"
    "‚Ä¢ PROCURE confrontantes em TODAS as dire√ß√µes (norte, sul, leste, oeste, nascente, poente, frente, fundos)\n"
    "‚Ä¢ SE MENOS DE 4 CONFRONTANTES: releia o texto procurando informa√ß√µes perdidas\n\n"

    "‚ö†Ô∏è CONSEQU√äNCIAS:\n"
    "‚ùå UM confrontante perdido = usucapi√£o pode ser NEGADO\n"
    "‚úÖ TODOS confrontantes identificados = processo bem fundamentado\n\n"

    "üìã AN√ÅLISE COMPLETA OBRIGAT√ìRIA:\n\n"

    "1Ô∏è‚É£ IDENTIFICA√á√ÉO DE MATR√çCULAS:\n"
    "‚Ä¢ Encontre todas as matr√≠culas presentes (n√∫meros, mesmo com varia√ß√µes de formata√ß√£o)\n"
    "‚Ä¢ Para cada matr√≠cula: extraia n√∫mero, LOTE, QUADRA, propriet√°rios ATUAIS, descri√ß√£o, confrontantes\n"
    "‚Ä¢ Ignore vendedores/doadores antigos - considere apenas √∫ltimos propriet√°rios\n"
    "‚Ä¢ Determine qual √© a matr√≠cula principal (objeto do usucapi√£o)\n\n"

    "2Ô∏è‚É£ AN√ÅLISE EXTREMAMENTE RIGOROSA DE CONFRONTANTES:\n"
    "üìç ONDE PROCURAR CONFRONTANTES:\n"
    "‚Ä¢ EXCLUSIVAMENTE na DESCRI√á√ÉO DA MATR√çCULA PRINCIPAL\n"
    "‚Ä¢ Se√ß√µes 'CONFRONTA√á√ïES', 'LIMITES', 'DIVISAS' da matr√≠cula principal\n"
    "‚Ä¢ N√ÉO buscar confrontantes em outros documentos ou matr√≠culas anexadas\n"
    "‚Ä¢ FOCO TOTAL: apenas a descri√ß√£o do im√≥vel da matr√≠cula objeto do usucapi√£o\n\n"

    "üîç PALAVRAS-CHAVE OBRIGAT√ìRIAS:\n"
    "‚Ä¢ 'confronta', 'limita', 'divisa', 'ao norte/sul/leste/oeste'\n"
    "‚Ä¢ 'frente', 'fundos', 'laterais', 'adjacente', 'vizinho'\n\n"

    "üéØ TIPOS DE CONFRONTANTES:\n"
    "‚Ä¢ LOTES: 'lote 11', 'lote n¬∫ 09' ‚Ä¢ MATR√çCULAS: 'matr√≠cula 1.234'\n"
    "‚Ä¢ PESSOAS: nomes completos ‚Ä¢ EMPRESAS: raz√µes sociais\n"
    "‚Ä¢ VIAS P√öBLICAS: ruas, avenidas (PROPRIEDADE DO MUNIC√çPIO)\n"
    "‚Ä¢ RODOVIAS ESTADUAIS: apenas estas s√£o de PROPRIEDADE DO ESTADO\n"
    "‚Ä¢ ENTES P√öBLICOS: Estado, Munic√≠pio ‚Ä¢ ACIDENTES GEOGR√ÅFICOS: rios, c√≥rregos\n\n"

    "‚ö° REGRAS CR√çTICAS:\n"
    "‚Ä¢ LEIA PALAVRA POR PALAVRA da descri√ß√£o do im√≥vel principal\n"
    "‚Ä¢ CONFRONTANTES: buscar SOMENTE na matr√≠cula principal, N√ÉO em outras matr√≠culas\n"
    "‚Ä¢ TODO lote tem 4 lados = m√≠nimo 4 confrontantes\n"
    "‚Ä¢ QUANDO MATR√çCULA N√ÉO ANEXADA: indique 'Matr√≠cula n√£o anexada' no campo matr√≠cula\n"
    "‚Ä¢ EXPRESSE CLARAMENTE quando confrontantes n√£o t√™m matr√≠cula anexada\n"
    "‚Ä¢ Se menos de 4: RELEIA procurando mais\n"
    "‚Ä¢ N√ÉO suponha, EXTRAIA exatamente como escrito\n\n"

    "3Ô∏è‚É£ CADEIA DOMINIAL COMPLETA:\n"
    "‚Ä¢ Analise hist√≥rico completo de propriet√°rios desde titula√ß√£o original\n"
    "‚Ä¢ Procure se√ß√µes: 'REGISTRO', 'TRANSMISS√ïES', 'AVERBA√á√ïES'\n"
    "‚Ä¢ Para cada transmiss√£o: data, tipo, propriet√°rio anterior, novo propriet√°rio, percentual, valor\n"
    "‚Ä¢ Co-propriedade: trate cada percentual como cadeia aut√¥noma\n\n"

    "4Ô∏è‚É£ RESTRI√á√ïES E GRAVAMES:\n"
    "‚Ä¢ Identifique restri√ß√µes n√£o baixadas: PENHORA, HIPOTECA, INDISPONIBILIDADE\n"
    "‚Ä¢ Verifique status: procure 'BAIXA', 'CANCELAMENTO', 'EXTIN√á√ÉO'\n"
    "‚Ä¢ ATEN√á√ÉO ESPECIAL: direitos do Estado de Mato Grosso do Sul\n\n"


    "üö® VERIFICA√á√ïES OBRIGAT√ìRIAS:\n"
    "‚Ä¢ Estado de MS como confrontante ou com direitos registrados?\n"
    "‚Ä¢ M√≠nimo 4 confrontantes identificados?\n"
    "‚Ä¢ Propriet√°rios atuais confirmados?\n"
    "‚Ä¢ Todas as matr√≠culas mapeadas?\n\n"

    "üî• ZERO TOLER√ÇNCIA para confrontantes perdidos. Cada um √© VITAL.\n\n"

    "Considere linguagem arcaica, abrevia√ß√µes, varia√ß√µes tipogr√°ficas e OCR imperfeito. "
    "Para an√°lise visual: leia todo texto vis√≠vel incluindo tabelas, carimbos e anota√ß√µes manuscritas."
)

# Instru√ß√µes espec√≠ficas por tipo de an√°lise
ANALYSIS_INSTRUCTIONS = {
    'aggregate': (
        "Voc√™ receber√° texto extra√≠do de documentos de um processo de usucapi√£o contendo m√∫ltiplas matr√≠culas. "
        "Aplique todas as instru√ß√µes do sistema para an√°lise completa.\n\n"
    ),
    'vision': (
        "Analise visualmente as imagens de matr√≠culas imobili√°rias. "
        "Leia todo o texto vis√≠vel (tabelas, carimbos, anota√ß√µes) considerando ru√≠dos de OCR. "
        "Aplique todas as instru√ß√µes do sistema com o mesmo rigor da an√°lise textual.\n\n"
    ),
    'partial': (
        "Voc√™ receber√° UM TRECHO de uma matr√≠cula. Retorne APENAS JSON com:\n"
        '{ "confrontantes": ["..."], "evidence": ["trecho literal..."] }\n'
        "Liste confrontantes exatamente como aparecem no trecho e evid√™ncias curtas.\n\n"
    )
}

# Esquema JSON padronizado
JSON_SCHEMA = '''
Responda em JSON com este esquema:
{
  "matriculas_encontradas": [
    {
      "numero": "12345",
      "lote": "10",
      "quadra": "21",
      "proprietarios": ["Nome 1", "Nome 2"],
      "descricao": "descri√ß√£o do im√≥vel",
      "confrontantes": ["lote 11", "confrontante 2"],
      "evidence": ["trecho literal 1", "trecho literal 2"],
      "cadeia_dominial": [
        {
          "data": "01/01/2020",
          "tipo_transmissao": "compra e venda",
          "proprietario_anterior": "Jo√£o Silva",
          "novo_proprietario": "Maria Santos",
          "percentual": "100%",
          "valor": "R$ 100.000,00",
          "registro": "R.1"
        }
      ],
      "restricoes": [
        {
          "tipo": "hipoteca",
          "data_registro": "15/06/2019",
          "credor": "Banco XYZ",
          "valor": "R$ 80.000,00",
          "situacao": "vigente",
          "data_baixa": null,
          "observacoes": "hipoteca para financiamento imobili√°rio"
        }
      ],
    }
  ],
  "matricula_principal": "12345",
  "matriculas_confrontantes": ["12346", "12347"],
  "lotes_confrontantes": [
    {
      "identificador": "lote 11",
      "tipo": "lote",
      "matricula_anexada": "12346",
      "direcao": "norte"
    },
    {
      "identificador": "lote 09",
      "tipo": "lote",
      "matricula_anexada": null,
      "direcao": "sul"
    },
    {
      "identificador": "Rua das Flores",
      "tipo": "via_publica",
      "matricula_anexada": null,
      "direcao": "leste"
    },
    {
      "identificador": "BR-163",
      "tipo": "rodovia_estadual",
      "matricula_anexada": null,
      "direcao": "oeste"
    }
  ],
  "matriculas_nao_confrontantes": ["12348"],
  "lotes_sem_matricula": ["lote 09"],
  "confrontacao_completa": true|false|null,
  "proprietarios_identificados": {"12345": ["Nome"], "12346": ["Nome2"]},
  "resumo_analise": {
    "cadeia_dominial_completa": {
      "12345": [
        {"proprietario": "Origem/Titula√ß√£o", "periodo": "at√© 2015", "percentual": "100%"},
        {"proprietario": "Jo√£o Silva", "periodo": "2015-2020", "percentual": "100%"},
        {"proprietario": "Maria Santos", "periodo": "2020-atual", "percentual": "100%"}
      ]
    },
    "restricoes_vigentes": [
      {"tipo": "hipoteca", "credor": "Banco XYZ", "valor": "R$ 80.000,00", "status": "vigente"}
    ],
    "restricoes_baixadas": [
      {"tipo": "penhora", "data_baixa": "10/12/2021", "motivo": "quita√ß√£o judicial"}
    ],
    "estado_ms_direitos": {
      "tem_direitos": true|false,
      "detalhes": [
        {"matricula": "12345", "tipo_direito": "credor_hipoteca", "status": "vigente", "valor": "R$ 50.000,00"}
      ],
      "criticidade": "alta|media|baixa",
      "observacao": "Estado de MS possui hipoteca vigente na matr√≠cula principal"
    }
  },
  "confidence": 0.0-1.0,
  "reasoning": "explica√ß√£o detalhada da an√°lise"
}

TIPOS DE CONFRONTANTES:
- 'lote': lotes numerados (ex: lote 11, lote 15)
- 'matricula': matr√≠culas identificadas por n√∫mero
- 'pessoa': nomes de pessoas propriet√°rias
- 'via_publica': ruas, avenidas, pra√ßas
- 'estado': Estado, Munic√≠pio, Uni√£o
- 'outros': c√≥rregos, rios, outros elementos
'''

def build_prompt(prompt_type: str) -> str:
    """Retorna o prompt unificado para o tipo informado.

    prompt_type: 'system', 'aggregate', 'vision' ou 'partial'
    """
    prompt = prompt_type.lower().strip()

    if prompt == 'system':
        return UNIFIED_SYSTEM_PROMPT

    if prompt in ANALYSIS_INSTRUCTIONS:
        if prompt == 'partial':
            return UNIFIED_SYSTEM_PROMPT + "\n\n" + ANALYSIS_INSTRUCTIONS[prompt]
        else:
            return UNIFIED_SYSTEM_PROMPT + "\n\n" + ANALYSIS_INSTRUCTIONS[prompt] + JSON_SCHEMA

    raise ValueError("prompt_type must be 'system', 'aggregate', 'vision', or 'partial'")

def build_analysis_prompt(mode: str) -> str:
    """Conveni√™ncia para obter prompt de an√°lise textual ou visual."""
    prompt = mode.lower().strip()
    if prompt == 'text':
        return build_prompt('aggregate')
    elif prompt == 'vision':
        return build_prompt('vision')
    else:
        raise ValueError("mode must be 'text' or 'vision'")


def build_full_report_prompt(data_json: str) -> str:
    """Monta prompt para solicitar um relat√≥rio textual completo √† LLM."""
    template = textwrap.dedent(
        f"""
<context_gathering>
        Voc√™ √© um assessor jur√≠dico especializado em usucapi√£o, auxiliando o Procurador do Estado de Mato Grosso do Sul em processo judicial no qual o Estado foi citado. 
        Sua tarefa √© redigir um **relat√≥rio t√©cnico completo, objetivo e fundamentado**, analisando exclusivamente o quadro de informa√ß√µes estruturadas fornecido.

        O relat√≥rio deve avaliar se o Estado de Mato Grosso do Sul possui interesse jur√≠dico no feito, considerando cadeia dominial, confronta√ß√µes, restri√ß√µes e direitos incidentes.
        </context_gathering>

        <structured_output>
        T√≠tulo inicial: **RELAT√ìRIO COMPLETO DO IM√ìVEL**

        Ordem obrigat√≥ria das se√ß√µes:
        1. **CONTEXTO** ‚Äì s√≠ntese da matr√≠cula principal, localiza√ß√£o (quadra, lote), propriet√°rios atuais e anteriores, cadeia dominial e informa√ß√µes gerais.  
        2. **CONFRONTA√á√ïES** ‚Äì an√°lise detalhada dos confrontantes, indicando quais possuem matr√≠cula identificada, quais n√£o possuem e as implica√ß√µes jur√≠dicas.  
        3. **DIREITOS E RESTRI√á√ïES** ‚Äì descri√ß√£o minuciosa de √¥nus, hipotecas, penhoras, direitos do Estado ou de terceiros e respectivos status (vigente, baixado etc.).  
        4. **AN√ÅLISE CR√çTICA** ‚Äì avalia√ß√£o fundamentada sobre consist√™ncia, sufici√™ncia e eventuais conflitos de informa√ß√£o.  
        5. **LACUNAS IDENTIFICADAS** ‚Äì listar dados ausentes ou insuficientes (ex.: confrontantes sem matr√≠cula, cadeias dominiais incompletas, restri√ß√µes n√£o detalhadas).  
        6. **RECOMENDA√á√ïES** ‚Äì indicar medidas necess√°rias (ex.: dilig√™ncias cartor√°rias, notifica√ß√µes a terceiros, pesquisa complementar).  
        7. **PARECER FINAL** ‚Äì concluir de forma direta se, diante dos elementos apresentados, h√° ou n√£o interesse jur√≠dico do Estado de Mato Grosso do Sul no processo de usucapi√£o, mencionando explicitamente as matr√≠culas, lotes e restri√ß√µes relevantes.

        </structured_output>

        <rules>
        - Responder **sempre em portugu√™s do Brasil**.  
        - N√£o utilizar sauda√ß√µes, frases introdut√≥rias gen√©ricas nem termos t√©cnicos de inform√°tica (como "JSON").  
        - Quando houver aus√™ncia de informa√ß√£o, escrever: ‚ÄúN√£o informado no quadro‚Äù e explicar a relev√¢ncia jur√≠dica da lacuna.  
        - Converter express√µes booleanas ou t√©cnicas (true/false/null) para linguagem jur√≠dica: ‚ÄúSim‚Äù, ‚ÄúN√£o‚Äù ou ‚ÄúN√£o informado‚Äù.  
        - Citar n√∫meros de matr√≠culas, lotes, propriet√°rios e confrontantes sempre que presentes.  
        - Nunca inventar ou presumir dados n√£o constantes no quadro.  
        </rules>

        <dados>
        QUADRO DE INFORMA√á√ïES ESTRUTURADAS:  
        <<IN√çCIO DOS DADOS>>  
        {data_json}  
        <<FIM DOS DADOS>>
        </dados>
    """
    )
    return template.strip()




# Compatibilidade com c√≥digo existente
SYSTEM_PROMPT = build_prompt('system')
AGGREGATE_PROMPT = build_analysis_prompt('text')
PARTIAL_PROMPT = build_prompt('partial')


def _safe_get_dict(data, key, default=None):
    """Retorna valor do dicion√°rio garantindo que seja do tipo correto."""
    if default is None:
        default = {}
    
    value = data.get(key, default)
    if not isinstance(value, dict):
        return default
    return value

def _safe_get_list(data, key, default=None):
    """Retorna valor do dicion√°rio garantindo que seja uma lista."""
    if default is None:
        default = []
    
    value = data.get(key, default)
    if not isinstance(value, list):
        return default
    return value

def _safe_process_matricula_data(m_data):
    """Processa dados de matr√≠cula de forma robusta, evitando erros com campos vazios."""
    if not isinstance(m_data, dict):
        return None
    
    try:
        # Processa cadeia dominial
        cadeia_dominial_obj = []
        cadeia_data = _safe_get_list(m_data, "cadeia_dominial")
        for transmissao_data in cadeia_data:
            if isinstance(transmissao_data, dict):
                transmissao = TransmissaoInfo(
                    data=transmissao_data.get("data"),
                    tipo_transmissao=transmissao_data.get("tipo_transmissao"),
                    proprietario_anterior=transmissao_data.get("proprietario_anterior"),
                    novo_proprietario=transmissao_data.get("novo_proprietario"),
                    percentual=transmissao_data.get("percentual"),
                    valor=transmissao_data.get("valor"),
                    registro=transmissao_data.get("registro")
                )
                cadeia_dominial_obj.append(transmissao)
        
        # Processa restri√ß√µes
        restricoes_obj = []
        restricoes_data = _safe_get_list(m_data, "restricoes")
        for restricao_data in restricoes_data:
            if isinstance(restricao_data, dict):
                restricao = RestricaoInfo(
                    tipo=restricao_data.get("tipo", ""),
                    data_registro=restricao_data.get("data_registro"),
                    credor=restricao_data.get("credor"),
                    valor=restricao_data.get("valor"),
                    situacao=restricao_data.get("situacao", "vigente"),
                    data_baixa=restricao_data.get("data_baixa"),
                    observacoes=restricao_data.get("observacoes")
                )
                restricoes_obj.append(restricao)
        
        
        # Processa listas principais com valida√ß√£o
        proprietarios = _safe_get_list(m_data, "proprietarios")
        confrontantes_list = _safe_get_list(m_data, "confrontantes")
        evidence = _safe_get_list(m_data, "evidence")
        
        matricula = MatriculaInfo(
            numero=str(m_data.get("numero", "")),
            proprietarios=proprietarios,
            descricao=str(m_data.get("descricao", "")),
            confrontantes=confrontantes_list,
            evidence=evidence,
            lote=m_data.get("lote"),
            quadra=m_data.get("quadra"),
            cadeia_dominial=cadeia_dominial_obj,
            restricoes=restricoes_obj,
        )
        return matricula
        
    except Exception as e:
        print(f"‚ö†Ô∏è Erro ao processar dados da matr√≠cula: {e}")
        return None

def analyze_with_vision_llm(model: str, file_path: str) -> AnalysisResult:
    """
    Analisa documento usando vis√£o computacional da LLM (an√°lise direta de imagens).
    """
    fname_placeholder = os.path.basename(file_path)
    
    try:
        print(f"üîç Convertendo {fname_placeholder} para an√°lise visual...")
        
        # Converte arquivo para imagens
        ext = os.path.splitext(file_path.lower())[1]
        if ext == ".pdf":
            # Verifica o n√∫mero de p√°ginas ANTES de processar
            try:
                total_pages = get_pdf_page_count(file_path)
                print(f"üìä PDF cont√©m {total_pages} p√°gina(s)")
            except Exception as e:
                print(f"‚ö†Ô∏è Erro ao contar p√°ginas: {e}")
                total_pages = 0
            
            # Removido limite de p√°ginas - processar√° qualquer quantidade
            if total_pages > 100:
                print(f"‚ö†Ô∏è PDF com {total_pages} p√°ginas - processamento pode demorar")
            
            try:
                images = pdf_to_images(file_path, max_pages=None)  # sem limite de p√°ginas
                print(f"üìÑ PDF convertido em {len(images) if images else 0} p√°gina(s)")
            except Exception as e:
                print(f"‚ùå Erro ao converter PDF: {e}")
                print(f"üîç Tipo do erro: {type(e).__name__}")
                raise ValueError(f"Erro ao converter PDF para imagens: {e}")
                
        elif ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"]:
            try:
                images = [Image.open(file_path)]
                print(f"üñºÔ∏è Imagem carregada para an√°lise")
            except Exception as e:
                print(f"‚ùå Erro ao abrir imagem: {e}")
                raise ValueError(f"Erro ao abrir imagem: {e}")
        else:
            raise ValueError(f"Formato de arquivo n√£o suportado para an√°lise visual: {ext}")
        
        if not images:
            raise ValueError("N√£o foi poss√≠vel extrair imagens do arquivo")
        
        # Valida√ß√£o das imagens
        print(f"üîç Validando {len(images)} imagem(ns)...")
        images_validas = []
        for i, img in enumerate(images):
            try:
                if img and hasattr(img, 'size') and img.size[0] > 0 and img.size[1] > 0:
                    images_validas.append(img)
                else:
                    print(f"‚ö†Ô∏è Imagem {i+1} inv√°lida ou vazia")
            except Exception as e:
                print(f"‚ö†Ô∏è Erro ao validar imagem {i+1}: {e}")
        
        if not images_validas:
            raise ValueError("Nenhuma imagem v√°lida foi extra√≠da do arquivo")
            
        images = images_validas
        print(f"‚úÖ {len(images)} imagem(ns) v√°lida(s) para processar")
        
        print(f"üîÑ Preparando {len(images)} imagem(ns) para envio √† IA...")
        
        # Converte imagens para base64
        images_b64 = []
        total_size_kb = 0
        
        for i, img in enumerate(images):
            try:
                if not img or not hasattr(img, 'size'):
                    print(f"‚ö†Ô∏è Imagem {i+1} inv√°lida - pulando")
                    continue
                    
                print(f"üìê Processando imagem {i+1}: {img.size[0]}x{img.size[1]} pixels")
                b64 = image_to_base64(img, max_size=1536)  # tamanho maior para documentos
                if b64:
                    size_kb = len(b64) // 1024
                    total_size_kb += size_kb
                    images_b64.append(b64)
                    print(f"‚úÖ Imagem {i+1} preparada ({size_kb:.1f}KB)")
                    print(f"üìä Total acumulado: {total_size_kb:.1f}KB")
                else:
                    print(f"‚ö†Ô∏è Falha ao processar imagem {i+1}")
            except Exception as e:
                print(f"‚ùå Erro ao processar imagem {i+1}: {e}")
                print(f"üîç Tipo do erro: {type(e).__name__}")
                continue
        
        print(f"üìà TOTAL: {len(images_b64)} imagens preparadas, {total_size_kb:.1f}KB")
        
        # Processamento inteligente baseado no tamanho real
        # Ajuste din√¢mico da qualidade baseado no n√∫mero de p√°ginas
        if len(images_b64) > 50:
            print(f"‚ö†Ô∏è Muitas p√°ginas ({len(images_b64)}) - otimizando qualidade automaticamente")
            # Reconverte com qualidade menor para muitas p√°ginas
            print(f"üîç DEBUG: Tentando otimizar {len(images)} imagens originais...")
            images_b64_temp = []
            try:
                for i, img in enumerate(images):
                    print(f"üîÑ Reprocessando imagem {i+1}/{len(images)} com qualidade reduzida...")
                    # Qualidade menor para documentos grandes
                    b64 = image_to_base64(img, max_size=800, jpeg_quality=50)
                    if b64:
                        images_b64_temp.append(b64)
                        print(f"‚úÖ Imagem {i+1} otimizada com sucesso")
                    else:
                        print(f"‚ö†Ô∏è Falha ao otimizar imagem {i+1}")
                images_b64 = images_b64_temp
                total_size_kb = sum(len(img) // 1024 for img in images_b64) if images_b64 else 0
                print(f"üìà AP√ìS OTIMIZA√á√ÉO: {len(images_b64)} imagens, {total_size_kb:.1f}KB")
            except Exception as e:
                print(f"‚ùå ERRO na otimiza√ß√£o de muitas p√°ginas: {e}")
                print(f"üîç Tipo do erro: {type(e).__name__}")
                raise
        elif total_size_kb / 1024 > 20:  # Se maior que 20MB, otimiza
            print(f"‚ö†Ô∏è Payload grande ({total_size_kb/1024:.1f}MB) - otimizando qualidade")
            print(f"üîç DEBUG: Tentando otimizar {len(images)} imagens originais...")
            images_b64_temp = []
            try:
                for i, img in enumerate(images):
                    print(f"üîÑ Reprocessando imagem {i+1}/{len(images)} para reduzir tamanho...")
                    b64 = image_to_base64(img, max_size=1024, jpeg_quality=60)
                    if b64:
                        images_b64_temp.append(b64)
                        print(f"‚úÖ Imagem {i+1} comprimida com sucesso")
                    else:
                        print(f"‚ö†Ô∏è Falha ao comprimir imagem {i+1}")
                images_b64 = images_b64_temp
                total_size_kb = sum(len(img) // 1024 for img in images_b64) if images_b64 else 0
                print(f"üìà AP√ìS OTIMIZA√á√ÉO: {len(images_b64)} imagens, {total_size_kb:.1f}KB")
            except Exception as e:
                print(f"‚ùå ERRO na otimiza√ß√£o de payload grande: {e}")
                print(f"üîç Tipo do erro: {type(e).__name__}")
                raise
        
        if not images_b64:
            raise ValueError("N√£o foi poss√≠vel converter nenhuma imagem para envio")
        
        # Prompt unificado para analise visual
        vision_prompt = build_analysis_prompt('vision')

        print(f"[Vision] Enviando {len(images_b64)} imagem(ns) para {model}...")
        print(f"[Vision] Tamanho do prompt: {len(vision_prompt)} chars")

        print("[Vision] DEBUG: Iniciando chamada da API...")

        data = call_openrouter_vision(
            model=model,
            system_prompt=SYSTEM_PROMPT,
            user_prompt=vision_prompt,
            images_base64=images_b64,
            temperature=0.0,
            max_tokens=100000  # Tokens otimizados para an√°lise eficiente
        )
        
        print(f"‚úÖ Resposta da API recebida com sucesso")
        print(f"üîç DEBUG: Iniciando processamento da resposta...")
        
        # Debug da estrutura da resposta
        print(f"üîç Estrutura da resposta:")
        print(f"  - choices: {len(data.get('choices', []))} elementos")
        if data.get('choices'):
            choice = data['choices'][0]
            print(f"  - finish_reason: {choice.get('finish_reason')}")
            print(f"  - message keys: {list(choice.get('message', {}).keys())}")
            content = choice.get('message', {}).get('content', '')
            print(f"  - content length: {len(content) if content else 0}")
            if content:
                print(f"  - content preview: {content[:200]}...")
            else:
                print(f"  - content is: {repr(content)}")
        
        # Acesso seguro ao conte√∫do da resposta
        try:
            if not data.get("choices") or len(data["choices"]) == 0:
                raise IndexError("Lista 'choices' vazia na resposta da API")
            
            choice = data["choices"][0]
            if not choice.get("message"):
                raise KeyError("Campo 'message' n√£o encontrado na resposta")
                
            content = choice["message"].get("content", "")
            
            print(f"üîç Content final: {len(content) if content else 0} chars")
            if content:
                print(f"üìù Primeiros 500 chars: {content[:500]}")
            else:
                print(f"‚ö†Ô∏è Content est√° vazio ou None!")
                
        except (IndexError, KeyError, TypeError) as e:
            print(f"‚ùå Erro ao acessar conte√∫do da resposta: {e}")
            print(f"üìä Estrutura da resposta: {list(data.keys()) if isinstance(data, dict) else type(data)}")
            raise RuntimeError(f"Estrutura de resposta inv√°lida da API: {e}")
        
        try:
            print(f"üîç DEBUG: Iniciando limpeza do JSON...")
            # Limpa marcadores de c√≥digo markdown se presentes
            clean_content = clean_json_response(content)
            print(f"üîß JSON limpo para parse: {clean_content[:100]}...")
            print(f"üîç DEBUG: JSON limpo tem {len(clean_content)} caracteres")
            
            print(f"üîç DEBUG: Tentando fazer json.loads()...")
            parsed = json.loads(clean_content)
            print(f"‚úÖ JSON parsed com sucesso! Tipo: {type(parsed)}")
            print(f"üîç Keys no parsed: {list(parsed.keys()) if isinstance(parsed, dict) else 'n√£o √© dict'}")
        except json.JSONDecodeError as e:
            print(f"‚ùå Erro ao fazer parse do JSON da vis√£o: {e}")
            print(f"üìÑ Conte√∫do completo da resposta:")
            print(content)
            parsed = {
                "matriculas_encontradas": [],
                "matricula_principal": None,
                "matriculas_confrontantes": [],
                "lotes_confrontantes": [],
                "matriculas_nao_confrontantes": [],
                "lotes_sem_matricula": [],
                "confrontacao_completa": None,
                "proprietarios_identificados": {},
                "confidence": None,
                "reasoning": f"Erro de parsing JSON da an√°lise visual: {content[:500]}..."
            }

        # Converte dados das matr√≠culas para objetos MatriculaInfo usando processamento seguro
        matriculas_obj = []
        for m_data in parsed.get("matriculas_encontradas", []):
            matricula = _safe_process_matricula_data(m_data)
            if matricula is not None:
                matriculas_obj.append(matricula)

        # Processa lotes confrontantes
        print(f"üîç DEBUG: Processando lotes confrontantes...")
        lotes_confrontantes_obj = []
        lotes_confrontantes_raw = parsed.get("lotes_confrontantes", [])
        print(f"üîç lotes_confrontantes encontrados: {len(lotes_confrontantes_raw)} itens")
        
        try:
            for i, lote_data in enumerate(lotes_confrontantes_raw):
                print(f"üîç Processando lote {i+1}: {type(lote_data)}")
                if isinstance(lote_data, dict):
                    lote_confronta = LoteConfronta(
                        identificador=lote_data.get("identificador", ""),
                        tipo=lote_data.get("tipo", "outros"),
                        matricula_anexada=lote_data.get("matricula_anexada"),
                        direcao=lote_data.get("direcao")
                    )
                    lotes_confrontantes_obj.append(lote_confronta)
                    print(f"‚úÖ Lote {i+1} processado com sucesso")
                else:
                    print(f"‚ö†Ô∏è Lote {i+1} n√£o √© dict: {lote_data}")
        except Exception as e:
            print(f"‚ùå ERRO ao processar lotes confrontantes: {e}")
            print(f"üîç Tipo do erro: {type(e).__name__}")
            raise

        # Processa resumo da an√°lise com tratamento seguro
        resumo_data = _safe_get_dict(parsed, "resumo_analise")
        
        # Processa direitos do Estado de MS
        estado_ms_data = _safe_get_dict(resumo_data, "estado_ms_direitos")
        estado_ms_direitos = EstadoMSDireitos(
            tem_direitos=bool(estado_ms_data.get("tem_direitos", False)),
            detalhes=_safe_get_list(estado_ms_data, "detalhes"),
            criticidade=str(estado_ms_data.get("criticidade", "baixa")),
            observacao=str(estado_ms_data.get("observacao", ""))
        )
        
        resumo_analise = ResumoAnalise(
            cadeia_dominial_completa=_safe_get_list(resumo_data, "cadeia_dominial_completa"),
            restricoes_vigentes=_safe_get_list(resumo_data, "restricoes_vigentes"),
            restricoes_baixadas=_safe_get_list(resumo_data, "restricoes_baixadas"),
            estado_ms_direitos=estado_ms_direitos
        )

        return AnalysisResult(
            arquivo=fname_placeholder,
            matriculas_encontradas=matriculas_obj,
            matricula_principal=parsed.get("matricula_principal"),
            matriculas_confrontantes=parsed.get("matriculas_confrontantes", []),
            lotes_confrontantes=lotes_confrontantes_obj,
            matriculas_nao_confrontantes=parsed.get("matriculas_nao_confrontantes", []),
            lotes_sem_matricula=parsed.get("lotes_sem_matricula", []),
            confrontacao_completa=parsed.get("confrontacao_completa"),
            proprietarios_identificados=parsed.get("proprietarios_identificados", {}),
            resumo_analise=resumo_analise,
            confidence=parsed.get("confidence"),
            reasoning=parsed.get("reasoning", ""),
            raw_json=parsed
        )
        
    except Exception as e:
        # CAPTURE O ERRO E MOSTRE LOGS DETALHADOS ANTES DE RETORNAR
        print(f"üö® CAPTURADO ERRO GERAL na an√°lise visual!")
        print(f"‚ùå Tipo do erro: {type(e).__name__}")
        print(f"‚ùå Mensagem do erro: {str(e)}")
        print(f"‚ùå Arquivo sendo processado: {fname_placeholder}")
        
        # Traceback detalhado
        import traceback
        print(f"üìç Traceback completo:")
        traceback.print_exc()
        
        # Se an√°lise visual falhar, retorna erro estruturado
        return AnalysisResult(
            arquivo=fname_placeholder,
            matriculas_encontradas=[],
            matricula_principal=None,
            matriculas_confrontantes=[],
            lotes_confrontantes=[],
            matriculas_nao_confrontantes=[],
            lotes_sem_matricula=[],
            confrontacao_completa=None,
            proprietarios_identificados={},
            confidence=None,
            reasoning=f"Erro na an√°lise visual: {str(e)}",
            raw_json={}
        )

# Fun√ß√£o analyze_text_with_llm removida - pipeline textual obsoleto
    """
    Estrat√©gia:
    - Se texto for curto: chamada √∫nica com prompt agregado.
    - Se texto for longo: faz passadas parciais para extrair confrontantes/evid√™ncias, deduplica,
      e faz uma chamada final curta passando o resumo + um trecho representativo do original.
    """
    fname_placeholder = "<arquivo>"
    text = full_text.strip()
    if not text:
        return AnalysisResult(
            arquivo=fname_placeholder,
            matriculas_encontradas=[],
            matricula_principal=None,
            matriculas_confrontantes=[],
            lotes_confrontantes=[],
            matriculas_nao_confrontantes=[],
            lotes_sem_matricula=[],
            confrontacao_completa=None,
            proprietarios_identificados={},
            confidence=None,
            reasoning="Texto vazio ap√≥s OCR.",
            raw_json={}
        )

    chunks = chunk_text(text, max_chars=18000)

    # Caso simples: uma chamada direta
    if len(chunks) == 1:
        data = call_openrouter(
            model=model,
            system_prompt=SYSTEM_PROMPT,
            user_prompt=AGGREGATE_PROMPT + "\n\nTEXTO:\n" + chunks[0],
            temperature=0.0,
            max_tokens=1500
        )
        # Acesso seguro ao conte√∫do
        try:
            if not data.get("choices") or len(data["choices"]) == 0:
                raise IndexError("Lista 'choices' vazia na resposta da API")
            content = data["choices"][0]["message"]["content"]
        except (IndexError, KeyError, TypeError) as e:
            print(f"Erro ao acessar conte√∫do da resposta: {e}")
            parsed = {
                "matriculas_encontradas": [],
                "matricula_principal": None,
                "matriculas_confrontantes": [],
                "lotes_confrontantes": [],
                "matriculas_nao_confrontantes": [],
                "lotes_sem_matricula": [],
                "confrontacao_completa": None,
                "proprietarios_identificados": {},
                "confidence": None,
                "reasoning": f"Erro de estrutura da resposta da API: {e}"
            }
        else:
            try:
                clean_content = clean_json_response(content)
                parsed = json.loads(clean_content)
            except json.JSONDecodeError as e:
                print(f"Erro ao fazer parse do JSON: {e}")
                print(f"Conte√∫do da resposta: {content[:500]}")
                parsed = {
                    "matriculas_encontradas": [],
                    "matricula_principal": None,
                    "matriculas_confrontantes": [],
                    "lotes_confrontantes": [],
                    "matriculas_nao_confrontantes": [],
                    "lotes_sem_matricula": [],
                    "confrontacao_completa": None,
                    "proprietarios_identificados": {},
                    "confidence": None,
                    "reasoning": f"Erro de parsing JSON: {content}"
                }

        # Converte dados das matr√≠culas para objetos MatriculaInfo usando processamento seguro
        matriculas_obj = []
        for m_data in parsed.get("matriculas_encontradas", []):
            matricula = _safe_process_matricula_data(m_data)
            if matricula is not None:
                matriculas_obj.append(matricula)

        return AnalysisResult(
            arquivo=fname_placeholder,
            matriculas_encontradas=matriculas_obj,
            matricula_principal=parsed.get("matricula_principal"),
            matriculas_confrontantes=parsed.get("matriculas_confrontantes", []),
            lotes_confrontantes=[],  # TODO: implementar processamento de lotes para texto
            matriculas_nao_confrontantes=parsed.get("matriculas_nao_confrontantes", []),
            lotes_sem_matricula=parsed.get("lotes_sem_matricula", []),
            confrontacao_completa=parsed.get("confrontacao_completa"),
            proprietarios_identificados=parsed.get("proprietarios_identificados", {}),
            confidence=parsed.get("confidence"),
            reasoning=parsed.get("reasoning", ""),
            raw_json=parsed
        )

    # Texto longo: pipeline parcial
    all_confrontantes = []
    all_evidence = []

    for i, ch in enumerate(chunks, 1):
        data = call_openrouter(
            model=model,
            system_prompt=SYSTEM_PROMPT,
            user_prompt=PARTIAL_PROMPT + "\n\nTRECHO:\n" + ch,
            temperature=0.0,
            max_tokens=800
        )
        # Acesso seguro ao conte√∫do do chunk
        try:
            if not data.get("choices") or len(data["choices"]) == 0:
                print(f"Chunk {i}: Lista 'choices' vazia")
                continue
            content = data["choices"][0]["message"]["content"]
        except (IndexError, KeyError, TypeError) as e:
            print(f"Erro no chunk {i} ao acessar resposta: {e}")
            continue
            
        try:
            clean_content = clean_json_response(content)
            parsed = json.loads(clean_content)
            if "confrontantes" in parsed and isinstance(parsed["confrontantes"], list):
                all_confrontantes.extend([c for c in parsed["confrontantes"] if isinstance(c, str)])
            if "evidence" in parsed and isinstance(parsed["evidence"], list):
                all_evidence.extend([e for e in parsed["evidence"] if isinstance(e, str)])
        except json.JSONDecodeError as e:
            print(f"Erro JSON no chunk {i}: {e}")
            print(f"Conte√∫do: {content[:200] if 'content' in locals() else 'N/A'}")
        except Exception as e:
            print(f"Erro no chunk {i}: {e}")
            # ignora erro parcial

    # Deduplica√ß√£o leve mantendo ordem
    seen = set()
    dedup_confrontantes = []
    for c in all_confrontantes:
        key = c.strip().lower()
        if key and key not in seen:
            seen.add(key)
            dedup_confrontantes.append(c.strip())

    # Chamada final curta com resumo
    resumo = {
        "confrontantes_coletados": dedup_confrontantes[:200],
        "amostras_evidencia": all_evidence[:20]
    }
    final_user = (
        AGGREGATE_PROMPT
        + "\n\nRESUMO PREPARADO (use para decidir):\n"
        + json.dumps(resumo, ensure_ascii=False, indent=2)
        + "\n\nNota: Decida com base no resumo; se faltar certeza, indique is_confrontante=null com reasoning."
    )

    data = call_openrouter(
        model=model,
        system_prompt=SYSTEM_PROMPT,
        user_prompt=final_user,
        temperature=0.0,
        max_tokens=900
    )
    # Acesso seguro ao conte√∫do da chamada final
    try:
        if not data.get("choices") or len(data["choices"]) == 0:
            raise IndexError("Lista 'choices' vazia na resposta final da API")
        content = data["choices"][0]["message"]["content"]
    except (IndexError, KeyError, TypeError) as e:
        print(f"Erro ao acessar resposta final: {e}")
        parsed = {
            "matriculas_encontradas": [],
            "matricula_principal": None,
            "matriculas_confrontantes": [],
            "lotes_confrontantes": [],
            "matriculas_nao_confrontantes": [],
            "lotes_sem_matricula": [],
            "confrontacao_completa": None,
            "proprietarios_identificados": {},
            "confidence": None,
            "reasoning": f"Erro de estrutura na resposta final da API: {e}"
        }
    else:
        try:
            clean_content = clean_json_response(content)
            parsed = json.loads(clean_content)
        except json.JSONDecodeError as e:
            print(f"Erro JSON na chamada final: {e}")
            print(f"Conte√∫do: {content[:500]}")
            parsed = {
                "matriculas_encontradas": [],
                "matricula_principal": None,
                "matriculas_confrontantes": [],
                "lotes_confrontantes": [],
                "matriculas_nao_confrontantes": [],
                "lotes_sem_matricula": [],
                "confrontacao_completa": None,
                "proprietarios_identificados": {},
                "confidence": None,
                "reasoning": f"Erro JSON: {content}"
            }

    # Converte dados das matr√≠culas para objetos MatriculaInfo
    matriculas_obj = []
    for m_data in parsed.get("matriculas_encontradas", []):
        if isinstance(m_data, dict):
            matricula = MatriculaInfo(
                numero=m_data.get("numero", ""),
                proprietarios=m_data.get("proprietarios", []),
                descricao=m_data.get("descricao", ""),
                confrontantes=m_data.get("confrontantes", []),
                evidence=m_data.get("evidence", [])
            )
            matriculas_obj.append(matricula)

    # Se n√£o encontrou matr√≠culas estruturadas, cria uma gen√©rica com dados coletados
    if not matriculas_obj and (dedup_confrontantes or all_evidence):
        matricula_generica = MatriculaInfo(
            numero="n√£o identificado",
            proprietarios=[],
            descricao="",
            confrontantes=dedup_confrontantes,
            evidence=all_evidence[:20],
            lote=None,
            quadra=None
        )
        matriculas_obj.append(matricula_generica)

    return AnalysisResult(
        arquivo=fname_placeholder,
        matriculas_encontradas=matriculas_obj,
        matricula_principal=parsed.get("matricula_principal"),
        matriculas_confrontantes=parsed.get("matriculas_confrontantes", []),
        lotes_confrontantes=[],
        matriculas_nao_confrontantes=parsed.get("matriculas_nao_confrontantes", []),
        lotes_sem_matricula=parsed.get("lotes_sem_matricula", []),
        confrontacao_completa=parsed.get("confrontacao_completa"),
        proprietarios_identificados=parsed.get("proprietarios_identificados", {}),
        confidence=parsed.get("confidence"),
        reasoning=parsed.get("reasoning", ""),
        raw_json=parsed
    )

# =========================
# Sistema de Feedback
# =========================
class FeedbackManager:
    """Gerencia o envio de feedback para Google Forms + Sheets"""
    
    def __init__(self):
        pass
        
    def solicitar_feedback(self, parent, dados_geracao):
        """Mostra dialog de feedback ap√≥s uma gera√ß√£o"""
        dialog = FeedbackDialog(parent, dados_geracao, self.enviar_feedback)
        
    def enviar_feedback(self, feedback_data):
        """Envia feedback para Google Forms"""
        thread = threading.Thread(
            target=self._enviar_feedback_async,
            args=(feedback_data,),
            daemon=True
        )
        thread.start()
        
    def _enviar_feedback_async(self, feedback_data):
        """Envio assincrono para nao travar a interface"""
        try:
            url = (GOOGLE_FORM_CONFIG.get("url") or "").strip()
            if not url:
                print("[Feedback] Google Forms nao configurado. Feedback descartado.")
                return

            field_map = GOOGLE_FORM_CONFIG.get("fields", {})
            form_data = {}

            for field_key in ("tipo", "descricao", "modelo", "timestamp", "versao"):
                field_id = field_map.get(field_key)
                value = feedback_data.get(field_key)
                if field_id and value is not None:
                    form_data[field_id] = value

            if not form_data:
                print("[Feedback] Nenhum campo valido configurado para envio. Feedback descartado.")
                return

            response = requests.post(
                url,
                data=form_data,
                timeout=10,
                headers={
                    "User-Agent": "FeedbackManager/1.0",
                    "Content-Type": "application/x-www-form-urlencoded"
                }
            )

            if response.status_code in (200, 302):
                print("[Feedback] Feedback enviado com sucesso.")
            else:
                print(f"[Feedback] Erro ao enviar feedback: {response.status_code}")

        except Exception as e:
            print(f"[Feedback] Erro de conexao: {e}")

class FeedbackDialog(tk.Toplevel):
    """Dialog para coleta de feedback do usu√°rio"""
    
    def __init__(self, parent, dados_geracao, callback_envio):
        super().__init__(parent)
        self.callback_envio = callback_envio
        self.dados_geracao = dados_geracao
        self.resultado = None
        
        self.configurar_janela()
        self.criar_interface()
        self.centralizar_janela()
        
    def configurar_janela(self):
        self.title("Feedback - Resultado da Gera√ß√£o")
        self.geometry("500x450")
        self.resizable(False, False)
        self.transient(self.master)
        self.grab_set()
        
    def criar_interface(self):
        # Frame principal
        main_frame = ttk.Frame(self, padding="20")
        main_frame.pack(fill="both", expand=True)
        
        # T√≠tulo
        titulo = ttk.Label(
            main_frame,
            text="O sistema identificou corretamente as confronta√ß√µes?",
            font=("Arial", 12, "bold")
        )
        titulo.pack(pady=(0, 20))
        
        # Op√ß√µes de resultado
        self.var_resultado = tk.StringVar(value="acertou")
        
        frame_opcoes = ttk.LabelFrame(main_frame, text="Resultado", padding="10")
        frame_opcoes.pack(fill="x", pady=(0, 20))
        
        ttk.Radiobutton(
            frame_opcoes,
            text="‚úÖ Acertou - As confronta√ß√µes est√£o corretas",
            variable=self.var_resultado,
            value="acertou"
        ).pack(anchor="w", pady=2)
        
        ttk.Radiobutton(
            frame_opcoes,
            text="‚ùå Errou - H√° problemas nas confronta√ß√µes",
            variable=self.var_resultado,
            value="errou"
        ).pack(anchor="w", pady=2)
        
        # Campo de descri√ß√£o do erro
        frame_descricao = ttk.LabelFrame(main_frame, text="Se errou, descreva o problema (opcional)", padding="10")
        frame_descricao.pack(fill="both", expand=True, pady=(0, 20))
        
        # Frame para texto e scrollbar
        text_frame = ttk.Frame(frame_descricao)
        text_frame.pack(fill="both", expand=True)
        
        self.txt_descricao = tk.Text(
            text_frame,
            height=6,
            wrap="word",
            font=("Arial", 10)
        )
        self.txt_descricao.pack(side="left", fill="both", expand=True)
        
        # Scrollbar para o texto
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.txt_descricao.yview)
        scrollbar.pack(side="right", fill="y")
        self.txt_descricao.config(yscrollcommand=scrollbar.set)
        
        # Bot√µes
        frame_botoes = ttk.Frame(main_frame)
        frame_botoes.pack(fill="x", pady=(10, 0))
        
        ttk.Button(
            frame_botoes,
            text="üì§ Enviar Feedback",
            command=self.enviar_feedback
        ).pack(side="right", padx=(10, 0))
        
        ttk.Button(
            frame_botoes,
            text="‚è≠Ô∏è Pular",
            command=self.pular_feedback
        ).pack(side="right")
        
    def centralizar_janela(self):
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (self.winfo_width() // 2)
        y = (self.winfo_screenheight() // 2) - (self.winfo_height() // 2)
        self.geometry(f"+{x}+{y}")
        
    def enviar_feedback(self):
        tipo_usuario = "SUCESSO_AUTO" if self.var_resultado.get() == "acertou" else "ERRO"
        descricao_usuario = self.txt_descricao.get("1.0", "end-1c").strip()

        if not descricao_usuario:
            descricao_usuario = (
                "Usuario confirmou que as confrontacoes estao corretas."
                if tipo_usuario == "SUCESSO_AUTO"
                else "Usuario indicou problemas nas confrontacoes, sem detalhes adicionais."
            )

        detalhes_tecnicos = {
            "arquivo_processado": self.dados_geracao.get("arquivo", ""),
            "confrontacoes_encontradas": self.dados_geracao.get("confrontacoes", 0),
            "tempo_processamento": self.dados_geracao.get("tempo", 0),
            "modelo_ia_usado": self.dados_geracao.get("modelo", DEFAULT_MODEL),
            "resultado_usuario": self.var_resultado.get()
        }

        resumo_detalhes = "\n".join(
            f"{chave}: {valor}"
            for chave, valor in detalhes_tecnicos.items()
            if valor not in (None, "")
        )

        descricao_completa = descricao_usuario
        if resumo_detalhes:
            descricao_completa = f"{descricao_usuario}\n\n[Detalhes tecnicos]\n{resumo_detalhes}"

        feedback_data = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "tipo": tipo_usuario,
            "descricao": descricao_completa,
            "processo": self.dados_geracao.get("processo")
                        or self.dados_geracao.get("arquivo")
                        or "N/A",
            "modelo": self.dados_geracao.get("modelo", DEFAULT_MODEL),
            "versao": APP_VERSION,
            "detalhes_tecnicos": detalhes_tecnicos
        }

        self.callback_envio(feedback_data)
        self.destroy()

    def pular_feedback(self):
        self.destroy()

# =========================
# GUI Principal
# =========================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("1100x700")
        self.minsize(1000, 650)

        self.files: List[str] = []
        self.results: Dict[str, AnalysisResult] = {}
        self.queue = queue.Queue()

        # Sistema de Feedback Inteligente
        self.feedback_system = initialize_feedback_system(
            app_version=APP_VERSION,
            modelo_llm=DEFAULT_MODEL
        )

        # Sistema de Auto-atualiza√ß√£o
        self.updater = create_updater()
        self.updater.auto_update = False
        self.updater.parent_window = self
        self._update_window = None

        # Cache do relat√≥rio completo
        self.cached_full_report_text: Optional[str] = None
        self.cached_full_report_payload: Optional[str] = None

        self.create_widgets()
        self.poll_queue()

        # Configura evento de fechamento para feedback
        self.protocol("WM_DELETE_WINDOW", self._on_closing)

        # Inicia verifica√ß√£o de atualiza√ß√µes ap√≥s 2 segundos
        self.after(2000, self.check_for_updates)

    def create_widgets(self):
        # Top frame (controles)
        top = ttk.Frame(self)
        top.pack(fill="x", padx=10, pady=8)

        self.btn_add = ttk.Button(top, text="Adicionar PDFs/Imagens", command=self.add_files)
        self.btn_add.pack(side="left")

        self.btn_remove = ttk.Button(top, text="Remover Selecionados", command=self.remove_selected)
        self.btn_remove.pack(side="left", padx=(8,0))

        ttk.Label(top, text="Matr√≠cula Principal (opcional):").pack(side="left", padx=(16,4))
        self.matricula_var = tk.StringVar()
        self.matricula_entry = ttk.Entry(top, textvariable=self.matricula_var, width=15)
        self.matricula_entry.pack(side="left")
        
        # Adiciona placeholder manual
        def add_placeholder():
            if not self.matricula_var.get():
                self.matricula_entry.insert(0, "ex: 12345")
                self.matricula_entry.config(foreground='gray')
        
        def remove_placeholder(event):
            if self.matricula_entry.get() == "ex: 12345":
                self.matricula_entry.delete(0, tk.END)
                self.matricula_entry.config(foreground='black')
        
        def validate_placeholder(event):
            if not self.matricula_entry.get():
                add_placeholder()
        
        self.matricula_entry.bind('<FocusIn>', remove_placeholder)
        self.matricula_entry.bind('<FocusOut>', validate_placeholder)
        add_placeholder()
        
        self.btn_process = ttk.Button(top, text="Processar", command=self.process_all)
        self.btn_process.pack(side="left", padx=12)

        self.btn_export = ttk.Button(top, text="Exportar CSV", command=self.export_csv)
        self.btn_export.pack(side="left")

        self.btn_feedback = ttk.Button(top, text="‚ö†Ô∏è Reportar Erro no Conte√∫do", command=self.reportar_erro_feedback, state="disabled")
        self.btn_feedback.pack(side="left", padx=(8,0))

        # Configura refer√™ncia no sistema de feedback
        self.feedback_system.set_feedback_button(self.btn_feedback)

        self.btn_update = ttk.Button(top, text="Verificar Atualiza√ß√µes", command=self.manual_check_updates)
        self.btn_update.pack(side="right", padx=(8,0))

        # Progress bar
        self.progress = ttk.Progressbar(top, orient="horizontal", mode="determinate", length=220)
        self.progress.pack(side="right")

        # Split: esquerda (lista arquivos) / direita (resultados)
        split = ttk.Panedwindow(self, orient="horizontal")
        split.pack(fill="both", expand=True, padx=10, pady=(0,10))

        # Esquerda: arquivos
        left = ttk.Frame(split)
        split.add(left, weight=1)

        ttk.Label(left, text="Arquivos anexados").pack(anchor="w", pady=(0,4))
        self.tree_files = ttk.Treeview(left, columns=("caminho",), show="headings", height=12)
        self.tree_files.heading("caminho", text="Caminho")
        self.tree_files.pack(fill="both", expand=True)
        self.tree_files.bind("<Delete>", lambda e: self.remove_selected())

        # Direita: resultados
        right = ttk.Frame(split)
        split.add(right, weight=2)

        # √Årea de alerta para direitos do Estado de MS
        alert_frame = ttk.Frame(right)
        alert_frame.pack(fill="x", padx=5, pady=(0,10))
        
        self.estado_alert_var = tk.StringVar()
        self.estado_alert_label = ttk.Label(
            alert_frame, 
            textvariable=self.estado_alert_var,
            font=("Arial", 10, "bold"),
            foreground="red",
            background="yellow",
            relief="solid",
            borderwidth=2,
            padding=5
        )
        # Label inicialmente oculto
        self.estado_alert_label.pack_forget()
        
        ttk.Label(right, text="Detalhamento das Matr√≠culas").pack(anchor="w", pady=(0,4))

        self.results_notebook = ttk.Notebook(right)
        self.results_notebook.pack(fill="both", expand=True)

        # Aba: matr√≠cula principal
        self.tab_principal = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.tab_principal, text="Matr√≠cula principal")

        principal_container = ttk.Frame(self.tab_principal)
        principal_container.pack(fill="both", expand=True, padx=5, pady=5)

        self.txt_principal = tk.Text(
            principal_container,
            wrap="word",
            height=15,
            state=tk.DISABLED,
            font=("Calibri", 11),
            padx=16,
            pady=12,
            spacing3=8,
            bg="#fafafa",
            relief="flat",
            borderwidth=1
        )
        principal_scroll = ttk.Scrollbar(principal_container, orient="vertical", command=self.txt_principal.yview)
        self.txt_principal.configure(yscrollcommand=principal_scroll.set)
        self.txt_principal.pack(side="left", fill="both", expand=True)
        principal_scroll.pack(side="right", fill="y")

        # Aba: confrontantes
        self.tab_confrontantes = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.tab_confrontantes, text="Confrontantes")

        lotes_frame = ttk.LabelFrame(self.tab_confrontantes, text="Lotes e matr√≠culas confrontantes")
        lotes_frame.pack(fill="both", expand=True, padx=5, pady=(5,2))

        self.tree_confrontantes_lotes = ttk.Treeview(lotes_frame, columns=("direcao", "identificador", "matricula", "proprietarios"), show="headings", height=8)
        self.tree_confrontantes_lotes.heading("direcao", text="Dire√ß√£o")
        self.tree_confrontantes_lotes.heading("identificador", text="Identificador")
        self.tree_confrontantes_lotes.heading("matricula", text="Matr√≠cula")
        self.tree_confrontantes_lotes.heading("proprietarios", text="Propriet√°rios")
        self.tree_confrontantes_lotes.column("direcao", width=90, anchor="center")
        self.tree_confrontantes_lotes.column("identificador", width=160, anchor="w")
        self.tree_confrontantes_lotes.column("matricula", width=110, anchor="center")
        self.tree_confrontantes_lotes.column("proprietarios", anchor="w")

        lotes_scroll = ttk.Scrollbar(lotes_frame, orient="vertical", command=self.tree_confrontantes_lotes.yview)
        self.tree_confrontantes_lotes.configure(yscrollcommand=lotes_scroll.set)
        self.tree_confrontantes_lotes.pack(side="left", fill="both", expand=True, padx=(0,2), pady=2)
        lotes_scroll.pack(side="right", fill="y", pady=2)

        outros_frame = ttk.LabelFrame(self.tab_confrontantes, text="Confrontantes especiais (vias, rios, Estado, etc.)")
        outros_frame.pack(fill="both", expand=True, padx=5, pady=(0,5))

        self.tree_confrontantes_outros = ttk.Treeview(outros_frame, columns=("tipo", "identificador", "direcao", "detalhes"), show="headings", height=6)
        self.tree_confrontantes_outros.heading("tipo", text="Tipo")
        self.tree_confrontantes_outros.heading("identificador", text="Identificador")
        self.tree_confrontantes_outros.heading("direcao", text="Dire√ß√£o")
        self.tree_confrontantes_outros.heading("detalhes", text="Observa√ß√µes / v√≠nculos")
        self.tree_confrontantes_outros.column("tipo", width=120, anchor="center")
        self.tree_confrontantes_outros.column("identificador", width=180, anchor="w")
        self.tree_confrontantes_outros.column("direcao", width=90, anchor="center")
        self.tree_confrontantes_outros.column("detalhes", anchor="w")

        outros_scroll = ttk.Scrollbar(outros_frame, orient="vertical", command=self.tree_confrontantes_outros.yview)
        self.tree_confrontantes_outros.configure(yscrollcommand=outros_scroll.set)
        self.tree_confrontantes_outros.pack(side="left", fill="both", expand=True, padx=(0,2), pady=2)
        outros_scroll.pack(side="right", fill="y", pady=2)

        # Aba: n√£o-confrontantes
        self.tab_nao_confrontantes = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.tab_nao_confrontantes, text="N√£o confrontantes")

        nao_conf_frame = ttk.Frame(self.tab_nao_confrontantes)
        nao_conf_frame.pack(fill="both", expand=True, padx=5, pady=5)

        self.tree_nao_confrontantes = ttk.Treeview(nao_conf_frame, columns=("matricula", "identificador", "proprietarios"), show="headings", height=12)
        self.tree_nao_confrontantes.heading("matricula", text="Matr√≠cula")
        self.tree_nao_confrontantes.heading("identificador", text="Identificador")
        self.tree_nao_confrontantes.heading("proprietarios", text="Propriet√°rios")
        self.tree_nao_confrontantes.column("matricula", width=120, anchor="center")
        self.tree_nao_confrontantes.column("identificador", width=200, anchor="w")
        self.tree_nao_confrontantes.column("proprietarios", anchor="w")

        nao_conf_scroll = ttk.Scrollbar(nao_conf_frame, orient="vertical", command=self.tree_nao_confrontantes.yview)
        self.tree_nao_confrontantes.configure(yscrollcommand=nao_conf_scroll.set)
        self.tree_nao_confrontantes.pack(side="left", fill="both", expand=True)
        nao_conf_scroll.pack(side="right", fill="y")
        # Bot√µes de a√ß√£o sobre resultado
        btns = ttk.Frame(right)
        btns.pack(fill="x", pady=6)
        self.btn_full_report = ttk.Button(btns, text="Gerar Relat√≥rio Completo", command=self.generate_full_report)
        self.btn_full_report.pack(side="left")

        # Campo de resumo (maior para melhor legibilidade)
        ttk.Label(right, text="Resumo da An√°lise").pack(anchor="w", pady=(10,4))
        self.txt_resumo = tk.Text(right, height=8, wrap=tk.WORD, state=tk.DISABLED, font=("TkDefaultFont", 10))
        self.txt_resumo.pack(fill="both", expand=True, pady=(0,6))

        # Log
        ttk.Label(self, text="Log / Mensagens").pack(anchor="w", padx=10)
        self.txt_log = tk.Text(self, height=8)
        self.txt_log.pack(fill="both", expand=False, padx=10, pady=(0,10))

    # ---------- A√ß√µes ----------
    def add_files(self):
        paths = filedialog.askopenfilenames(
            title="Selecione PDFs ou imagens",
            filetypes=[
                ("PDF e Imagem", "*.pdf;*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp;*.webp"),
                ("PDF", "*.pdf"),
                ("Imagens", "*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp;*.webp"),
                ("Todos", "*.*"),
            ]
        )
        if not paths:
            return
        if self.files:
            # Notifica sistema de feedback sobre sucesso impl√≠cito antes de limpar
            if hasattr(self, 'feedback_system'):
                # Se havia processo anterior sem feedback negativo, envia feedback positivo
                # (usu√°rio est√° adicionando novos arquivos = satisfeito com resultado anterior)
                self.feedback_system.on_relatorio_sucesso("novo_processamento")

            self.files.clear()
            for item in self.tree_files.get_children():
                self.tree_files.delete(item)
            self.log("Arquivos anteriores removidos.")
            self.cached_full_report_text = None
            self.cached_full_report_payload = None
        new = 0
        for p in paths:
            p = os.path.abspath(p)
            if p not in self.files:
                self.files.append(p)
                self.tree_files.insert("", "end", values=(p,))
                new += 1
        if new:
            self.log(f"{new} arquivo(s) adicionado(s).")

    def remove_selected(self):
        sel = self.tree_files.selection()
        removed = 0
        for item in sel:
            path = self.tree_files.item(item, "values")[0]
            if path in self.files:
                self.files.remove(path)
            self.tree_files.delete(item)
            removed += 1
        if removed:
            self.log(f"{removed} arquivo(s) removido(s).")

    def process_all(self):
        if not self.files:
            messagebox.showwarning("Nada a processar", "Adicione pelo menos um arquivo.")
            return
        model = DEFAULT_MODEL

        self.progress["value"] = 0
        self.progress["maximum"] = len(self.files)
        self.results.clear()
        self.clear_result_views()
        self.cached_full_report_text = None
        self.cached_full_report_payload = None

        t = threading.Thread(target=self._worker_process, args=(model,), daemon=True)
        t.start()

    def _worker_process(self, model: str):
        for idx, path in enumerate(self.files, 1):
            filename = os.path.basename(path)
            try:
                self.queue.put(("log", f"üìÑ Processando {filename} ({idx}/{len(self.files)})"))
                
                # Verifica se o arquivo existe e diagnostica problemas
                if not os.path.exists(path):
                    self.queue.put(("log", f"‚ùå Arquivo n√£o encontrado: {filename}"))
                    continue
                
                # Diagn√≥stico do arquivo
                diagnostico = self.diagnose_file_issues(path)
                if "‚ùå" in diagnostico or "‚ö†Ô∏è" in diagnostico:
                    self.queue.put(("log", f"üîç Diagn√≥stico: {diagnostico}"))
                
                # An√°lise visual direta com IA
                self.queue.put(("log", f"üëÅÔ∏è Analisando documento visualmente com IA..."))
                
                # Adiciona n√∫mero da matr√≠cula informado pelo usu√°rio (se houver)
                matricula_informada = self.matricula_var.get().strip()
                if matricula_informada and matricula_informada != "ex: 12345":
                    matricula_normalizada = matricula_informada.replace(".", "").replace(" ", "")
                    self.queue.put(("log", f"üìù Matr√≠cula de refer√™ncia informada: {matricula_normalizada}"))
                
                res = analyze_with_vision_llm(model, path)
                res.arquivo = filename
                self.results[path] = res

                # Log dos resultados principais
                if res.reasoning and "Erro na an√°lise visual" in res.reasoning:
                    # Extrai mais detalhes do erro com prote√ß√£o
                    try:
                        partes = res.reasoning.split("Erro na an√°lise visual: ")
                        if len(partes) > 1:
                            erro_detalhes = partes[-1][:100]
                        else:
                            erro_detalhes = res.reasoning[:100]
                    except Exception:
                        erro_detalhes = "Erro desconhecido"
                    
                    self.queue.put(("log", f"‚ö†Ô∏è Problema na an√°lise visual: {erro_detalhes}"))
                    self.queue.put(("log", f"üí° Poss√≠veis causas: arquivo muito grande, ileg√≠vel ou formato n√£o suportado"))
                elif res.matriculas_encontradas:
                    self.queue.put(("log", f"üìã {len(res.matriculas_encontradas)} matr√≠cula(s) identificada(s) visualmente"))
                    if res.matricula_principal:
                        self.queue.put(("log", f"üè† Matr√≠cula principal: {res.matricula_principal}"))
                    if res.matriculas_confrontantes:
                        self.queue.put(("log", f"üîó {len(res.matriculas_confrontantes)} matr√≠cula(s) confrontante(s)"))
                    if res.is_confrontante:
                        self.queue.put(("log", f"üèõÔ∏è Estado de MS identificado como confrontante"))
                else:
                    self.queue.put(("log", f"‚ö†Ô∏è Nenhuma matr√≠cula foi identificada no documento"))
                
                # Prepara dados para exibi√ß√£o na tabela
                mat_principal = res.matricula_principal or "N√£o identificada"
                mat_confrontantes = ", ".join(res.matriculas_confrontantes[:3]) + ("..." if len(res.matriculas_confrontantes) > 3 else "")
                if not mat_confrontantes:
                    mat_confrontantes = "Nenhuma"
                
                estado_ms = "SIM" if res.is_confrontante else "N√ÉO"
                
                
                # Formata confian√ßa (j√° vem como percentual da API)
                if res.confidence is not None:
                    confianca_pct = f"{int(res.confidence)}%"
                else:
                    confianca_pct = "N/A"
                
                # Resumo dos propriet√°rios
                proprietarios_resumo = []
                for numero, props in list(res.proprietarios_identificados.items())[:2]:
                    if props:
                        proprietarios_resumo.append(f"{numero}: {props[0]}" + ("..." if len(props) > 1 else ""))
                proprietarios_texto = " | ".join(proprietarios_resumo)
                if not proprietarios_texto:
                    proprietarios_texto = "N√£o identificados"

                # Mensagem de conclus√£o baseada no status
                if res.reasoning and "Erro na an√°lise visual" in res.reasoning:
                    self.queue.put(("log", f"‚ö†Ô∏è An√°lise de {filename} conclu√≠da com problemas (confian√ßa: {confianca_pct})"))
                elif res.matriculas_encontradas:
                    self.queue.put(("log", f"‚úÖ An√°lise de {filename} conclu√≠da com sucesso (confian√ßa: {confianca_pct})"))
                else:
                    self.queue.put(("log", f"‚ÑπÔ∏è An√°lise de {filename} conclu√≠da - nenhuma matr√≠cula identificada (confian√ßa: {confianca_pct})"))
                
                self.queue.put(("result", (path, res)))
                
            except Exception as e:
                error_msg = str(e)
                if "p√°ginas excede o limite m√°ximo" in error_msg:
                    self.queue.put(("log", f"üö´ {filename}: {error_msg}"))
                else:
                    self.queue.put(("log", f"‚ùå Erro ao processar {filename}: {error_msg}"))
            finally:
                self.queue.put(("progress", 1))

    def export_csv(self):
        if not self.results:
            messagebox.showinfo("Sem resultados", "Nada para exportar ainda.")
            return
        out = filedialog.asksaveasfilename(
            title="Salvar CSV",
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")]
        )
        if not out:
            return
        try:
            import csv
            with open(out, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f, delimiter=";")
                w.writerow(["arquivo", "matricula_principal", "matriculas_confrontantes", "estado_ms_confrontante", 
                           "confianca_percentual", "proprietarios", "reasoning"])
                for path, res in self.results.items():
                    # Formata confian√ßa (j√° vem como percentual da API)
                    confianca_pct = f"{int(res.confidence)}%" if res.confidence is not None else "N/A"
                    
                    # Prepara propriet√°rios para CSV
                    proprietarios_csv = []
                    for numero, props in res.proprietarios_identificados.items():
                        if props:
                            proprietarios_csv.append(f"{numero}: {'; '.join(props)}")
                    
                    w.writerow([
                        res.arquivo,
                        res.matricula_principal or "N√£o identificada",
                        " | ".join(res.matriculas_confrontantes),
                        "SIM" if res.is_confrontante else "N√ÉO",
                        confianca_pct,
                        " | ".join(proprietarios_csv),
                        res.reasoning.replace("\n", " ").strip()
                    ])
            self.log(f"CSV salvo em: {out}")
        except Exception as e:
            messagebox.showerror("Erro ao salvar", str(e))

    def poll_queue(self):
        try:
            while True:
                kind, payload = self.queue.get_nowait()
                if kind == "log":
                    self.log(payload)
                elif kind == "result":
                    # payload agora cont√©m: path, result_object (AnalysisResult)
                    path, result = payload
                    self.populate_results_tree(result)
                    self.update_summary(result)
                    # Atualiza alerta sobre direitos do Estado de MS
                    self.update_estado_alert()

                    # Notifica sistema de feedback sobre sucesso
                    numero_processo = result.numero_processo if hasattr(result, 'numero_processo') and result.numero_processo else os.path.basename(path)
                    self.feedback_system.on_relatorio_sucesso(numero_processo)
                elif kind == "progress":
                    val = self.progress["value"] + payload
                    self.progress["value"] = val
                    # Verifica se o processamento foi conclu√≠do
                    if val >= self.progress["maximum"] and val > 0:
                        pass  # Feedback autom√°tico j√° gerenciado pelo sistema inteligente
        except queue.Empty:
            pass
        self.after(100, self.poll_queue)

    def solicitar_feedback_processamento(self):
        """Solicita feedback ap√≥s completar o processamento de todos os arquivos"""
        if not self.results:
            return  # Sem resultados para avaliar
            
        # Coleta dados sobre o processamento
        total_matriculas = sum(len(result.matriculas_encontradas) for result in self.results.values())
        arquivos_processados = len(self.results)
        confrontacoes_identificadas = sum(len(result.matriculas_confrontantes) for result in self.results.values())
        
        dados_geracao = {
            "arquivo": f"{arquivos_processados} arquivo(s) processado(s)",
            "confrontacoes": confrontacoes_identificadas,
            "tempo": 0,  # Tempo ser√° calculado posteriormente se necess√°rio
            "modelo": DEFAULT_MODEL,
            "matriculas_encontradas": total_matriculas
        }
        
        # Agenda feedback para depois da interface ser atualizada
        pass  # Feedback autom√°tico gerenciado pelo sistema inteligente

    def check_for_updates(self):
        """Verifica se ha atualizacao e consulta o usuario antes de aplicar"""
        def update_thread():
            try:
                # Primeiro, sincroniza a vers√£o local com o GitHub
                self.updater.sync_version_with_github()

                # Depois verifica se h√° atualiza√ß√µes execut√°veis
                update_info = self.updater.check_for_updates()
                if update_info:
                    self.after(0, lambda: self._show_update_dialog(update_info, self.updater, automatic=True))
            except Exception:
                pass

        # Executa em thread separada para nao bloquear a interface
        threading.Thread(target=update_thread, daemon=True).start()

    def manual_check_updates(self):
        """Verifica√ß√£o manual de atualiza√ß√µes com feedback ao usu√°rio"""
        def update_thread():
            try:
                print("üîÑ Verificando atualiza√ß√µes manualmente...")

                # Cria updater com silent=False para debug
                debug_updater = create_updater()
                debug_updater.silent = False
                debug_updater.auto_update = False
                debug_updater.parent_window = self

                update_info = debug_updater.check_for_updates()
                print(f"üìã Resultado da verifica√ß√£o: {update_info}")

                if update_info:
                    self.after(0, lambda: self._show_update_dialog(update_info, debug_updater))
                else:
                    self.after(0, lambda: messagebox.showinfo(
                        "Atualizado",
                        "Voc√™ j√° est√° usando a vers√£o mais recente!",
                        parent=self
                    ))
            except Exception as e:
                print(f"‚ùå Erro na verifica√ß√£o: {e}")
                import traceback
                traceback.print_exc()
                self.after(0, lambda: messagebox.showerror(
                    "Erro de Atualiza√ß√£o",
                    f"Erro ao verificar atualiza√ß√µes: {e}",
                    parent=self
                ))

        threading.Thread(target=update_thread, daemon=True).start()

    def _show_update_dialog(self, update_info, updater, automatic=False):
        """Exibe janela dedicada para confirmar atualizacao"""
        existing = getattr(self, "_update_window", None)
        if existing is not None and existing.winfo_exists():
            try:
                existing.lift()
                existing.focus_force()
            except tk.TclError:
                self._update_window = None
            else:
                return

        window = tk.Toplevel(self)
        window.title("Atualizacao disponivel")
        window.geometry("480x420")
        window.minsize(420, 360)
        window.transient(self)
        window.grab_set()

        self._update_window = window

        def close_window():
            if self._update_window is window:
                self._update_window = None
            window.destroy()

        window.protocol("WM_DELETE_WINDOW", close_window)

        container = ttk.Frame(window)
        container.pack(fill="both", expand=True, padx=16, pady=14)

        current_version = getattr(updater, "current_version", None) or "-"
        header_text = f"Nova versao {update_info['version']} disponivel"
        ttk.Label(container, text=header_text, font=("Arial", 13, "bold")).pack(anchor="w")

        if current_version and current_version != "-":
            ttk.Label(container, text=f"Versao instalada: {current_version}", font=("Arial", 10)).pack(anchor="w", pady=(4, 0))

        origin_message = "Este aviso foi aberto automaticamente." if automatic else "Este aviso foi aberto manualmente."
        ttk.Label(container, text=origin_message, font=("Arial", 9), foreground="gray").pack(anchor="w", pady=(2, 10))

        ttk.Separator(container).pack(fill="x", pady=(0, 10))

        notes = (update_info.get("release_notes") or "").strip()
        notes_frame = ttk.LabelFrame(container, text="Notas da versao")
        notes_frame.pack(fill="both", expand=True)

        notes_text = tk.Text(notes_frame, wrap="word", height=8, font=("Consolas", 9))
        notes_text.pack(side="left", fill="both", expand=True, padx=0, pady=4)
        scrollbar = ttk.Scrollbar(notes_frame, orient="vertical", command=notes_text.yview)
        scrollbar.pack(side="right", fill="y", padx=0, pady=4)
        notes_text.configure(yscrollcommand=scrollbar.set)
        notes_text.insert("1.0", notes if notes else "Nenhuma nota de versao foi publicada.")
        notes_text.configure(state="disabled")

        ttk.Separator(container).pack(fill="x", pady=(12, 10))

        status_var = tk.StringVar(value='Clique em "Atualizar agora" para iniciar o processo.')
        ttk.Label(container, textvariable=status_var, wraplength=420, font=("Arial", 10)).pack(fill="x")

        progress_var = tk.IntVar(value=0)
        progress = ttk.Progressbar(container, maximum=100, variable=progress_var)
        progress.pack(fill="x", pady=(6, 6))

        buttons = ttk.Frame(container)
        buttons.pack(fill="x", pady=(10, 0))

        def start_update():
            btn_update.config(state="disabled")
            btn_later.config(state="disabled")
            status_var.set("Iniciando atualizacao...")
            progress_var.set(5)

            def run_update():
                previous_auto = getattr(updater, "auto_update", True)
                try:
                    updater.auto_update = True

                    def progress_callback(state, percent):
                        def update_ui():
                            status_var.set(state)
                            try:
                                progress_value = int(percent)
                            except Exception:
                                progress_value = 0
                            progress_var.set(max(0, min(100, progress_value)))
                        self.after(0, update_ui)

                    success = updater.update_if_available(progress_callback=progress_callback)
                    if not success:
                        def finish_no_update():
                            status_var.set("Nenhuma atualizacao foi aplicada.")
                            btn_update.config(state="normal")
                            btn_later.config(state="normal")
                        self.after(0, finish_no_update)
                except Exception as exc:
                    def handle_error():
                        status_var.set(f"Erro durante a atualizacao: {exc}")
                        btn_update.config(state="normal")
                        btn_later.config(state="normal")
                    self.after(0, handle_error)
                finally:
                    updater.auto_update = previous_auto

            threading.Thread(target=run_update, daemon=True).start()

        btn_update = ttk.Button(buttons, text="Atualizar agora", command=start_update)
        btn_update.pack(side="left")

        btn_later = ttk.Button(buttons, text="Mais tarde", command=close_window)
        btn_later.pack(side="right")

        window.focus_set()

    def reportar_erro_feedback(self):
        """Abre di√°logo para reportar erro no conte√∫do"""
        self.feedback_system.on_reportar_erro_manual(parent_window=self)

    def _on_closing(self):
        """M√©todo chamado ao fechar a aplica√ß√£o - envia feedback autom√°tico se necess√°rio"""
        self.feedback_system.on_fechamento_aplicacao()
        self.destroy()

    def clear_result_views(self):
        """Limpa os componentes utilizados para exibir os resultados."""
        self._set_principal_content_markdown("")
        trees = [
            getattr(self, "tree_confrontantes_lotes", None),
            getattr(self, "tree_confrontantes_outros", None),
            getattr(self, "tree_nao_confrontantes", None),
        ]
        for tree in trees:
            if tree is None:
                continue
            for item in tree.get_children():
                tree.delete(item)

    def _set_principal_content(self, text: str):
        """Atualiza o conte√∫do textual da aba de matr√≠cula principal."""
        widget = getattr(self, "txt_principal", None)
        if widget is None:
            return
        widget.configure(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        if text:
            widget.insert("1.0", text)
        widget.configure(state=tk.DISABLED)

    def _set_principal_content_markdown(self, markdown_text: str):
        """Atualiza o conte√∫do da aba de matr√≠cula principal com formata√ß√£o markdown."""
        widget = getattr(self, "txt_principal", None)
        if widget is None:
            return
        self._render_markdown_content(widget, markdown_text)
        widget.configure(state=tk.DISABLED)

    def _insert_placeholder_row(self, tree, message):
        """Insere uma linha informativa quando n√£o h√° dados dispon√≠veis."""
        if tree is None:
            return
        columns = tree["columns"]
        if isinstance(columns, str):
            columns = (columns,)
        num_cols = len(columns)
        if num_cols == 0:
            tree.insert("", "end", text=message)
            return
        values = []
        for idx in range(num_cols):
            values.append(message if idx == 0 else "")
        tree.insert("", "end", values=tuple(values))

    def populate_results_tree(self, result):
        """Atualiza as abas de resultados respeitando as especificidades de cada tipo de matr√≠cula."""
        self.clear_result_views()

        if not result:
            self._set_principal_content_markdown("# ‚ö†Ô∏è Nenhum resultado processado\n\nN√£o foi poss√≠vel processar o arquivo selecionado.")
            self._insert_placeholder_row(self.tree_confrontantes_lotes, "Sem confrontantes identificados.")
            self._insert_placeholder_row(self.tree_confrontantes_outros, "Sem confrontantes especiais.")
            self._insert_placeholder_row(self.tree_nao_confrontantes, "Sem matr√≠culas anexadas.")
            return

        matriculas_map = {mat.numero: mat for mat in result.matriculas_encontradas}

        def owners_for(matricula_num: Optional[str]) -> List[str]:
            if not matricula_num:
                return []
            proprietarios = result.proprietarios_identificados.get(matricula_num)
            if proprietarios:
                return [p for p in proprietarios if p and p.upper() != "N/A"]
            mat_obj = matriculas_map.get(matricula_num)
            if mat_obj and mat_obj.proprietarios:
                return [p for p in mat_obj.proprietarios if p and p.upper() != "N/A"]
            return []

        def format_direction(direction: Optional[str]) -> str:
            if not direction:
                return "N√£o informada"
            normalized = direction.strip()
            if not normalized:
                return "N√£o informada"
            if len(normalized) <= 3:
                return normalized.upper()
            return normalized[0].upper() + normalized[1:].lower()

        def join_with_overflow(items: List[str], limit: int = 3) -> str:
            items = [item for item in items if item]
            if not items:
                return ""
            if len(items) <= limit:
                return "; ".join(items)
            return "; ".join(items[:limit]) + f" (+{len(items) - limit})"

        matricula_principal_obj = matriculas_map.get(result.matricula_principal) if result.matricula_principal else None
        numero_principal = result.matricula_principal or "N√£o identificada"

        if result.confidence is None:
            confianca_display = "N/A"
        elif isinstance(result.confidence, (int, float)):
            if 0 <= result.confidence <= 1:
                confianca_display = f"{int(result.confidence * 100)}%"
            else:
                confianca_display = f"{int(result.confidence)}%"
        else:
            confianca_display = str(result.confidence)

        estado_ms_display = "SIM" if result.is_confrontante else "N√ÉO"
        lotes_confrontantes_lista = result.lotes_confrontantes or []
        matriculas_confrontantes_lista = result.matriculas_confrontantes or []
        matriculas_nao_confrontantes_lista = result.matriculas_nao_confrontantes or []
        confrontantes_total = len(lotes_confrontantes_lista) if lotes_confrontantes_lista else len(matriculas_confrontantes_lista)
        especiais_total = len([c for c in lotes_confrontantes_lista if (c.tipo or "").lower() not in {"lote", "matricula"}])
        missing_confrontantes = []
        for ident in result.lotes_sem_matricula or []:
            if ident and ident not in missing_confrontantes:
                missing_confrontantes.append(ident)
        for conf in lotes_confrontantes_lista:
            tipo_key = (conf.tipo or "").lower()
            if tipo_key in {"lote", "matricula"} and not conf.matricula_anexada:
                ident = (conf.identificador or "Confrontante sem identifica√ß√£o").strip()
                if ident and ident not in missing_confrontantes:
                    missing_confrontantes.append(ident)
        nao_confrontantes_total = len(matriculas_nao_confrontantes_lista)

        if missing_confrontantes:
            estado_ms_display = "DESCONHECIDO"
            estado_ms_detail = f"{len(missing_confrontantes)} confrontante(s) sem matr√≠cula anexada"
        else:
            estado_ms_display = "SIM" if result.is_confrontante else "N√ÉO"
            estado_ms_detail = ""

        lote_parts = []
        if matricula_principal_obj and matricula_principal_obj.lote:
            lote_parts.append(f"Lote {matricula_principal_obj.lote}")
        if matricula_principal_obj and matricula_principal_obj.quadra:
            lote_parts.append(f"Quadra {matricula_principal_obj.quadra}")
        estado_value = estado_ms_display if not estado_ms_detail else f"{estado_ms_display} ‚Äî {estado_ms_detail}"

        principal_lines: List[str] = []
        principal_lines.append("# Relat√≥rio da Matr√≠cula Principal")
        principal_lines.append("")

        # Se√ß√£o de Identifica√ß√£o
        principal_lines.append("## üìã Identifica√ß√£o")
        principal_lines.append(f"**Matr√≠cula:** {numero_principal}")
        principal_lines.append(f"**Arquivo analisado:** {result.arquivo}")
        if lote_parts:
            principal_lines.append(f"**Lote / Quadra:** {' / '.join(lote_parts)}")
        principal_lines.append("")

        # Se√ß√£o de An√°lise
        principal_lines.append("## üîç An√°lise Realizada")
        principal_lines.append(f"**Estado MS confrontante:** {estado_value}")
        principal_lines.append(f"**Confian√ßa da IA:** {confianca_display}")
        if result.confrontacao_completa is not None:
            status = "‚úÖ Sim" if result.confrontacao_completa else "‚ùå N√£o"
            principal_lines.append(f"**Confronta√ß√£o completa:** {status}")
        principal_lines.append("")

        # Se√ß√£o de Estat√≠sticas
        principal_lines.append("## üìä Estat√≠sticas")
        principal_lines.append(f"‚Ä¢ **Total de confrontantes:** {confrontantes_total}")
        if especiais_total:
            principal_lines.append(f"‚Ä¢ **Confrontantes especiais:** {especiais_total}")
        if missing_confrontantes:
            missing_display = join_with_overflow(missing_confrontantes, limit=3) or "Identificadores n√£o informados"
            principal_lines.append(f"‚Ä¢ **Sem matr√≠cula analis√°vel:** {missing_display}")
        principal_lines.append(f"‚Ä¢ **N√£o confrontantes anexados:** {nao_confrontantes_total}")
        principal_lines.append(f"‚Ä¢ **Matr√≠culas analisadas:** {len(result.matriculas_encontradas)}")
        principal_lines.append("")

        # Se√ß√£o de Propriet√°rios
        proprietarios_principal = owners_for(result.matricula_principal)
        principal_lines.append("## üë• Propriet√°rios Identificados")
        if proprietarios_principal:
            for owner in proprietarios_principal:
                principal_lines.append(f"‚Ä¢ **{owner}**")
        else:
            principal_lines.append("‚Ä¢ *N√£o informado*")

        self._set_principal_content_markdown("\n".join(principal_lines))

        lotes_normais = []
        especiais = []
        for conf in lotes_confrontantes_lista:
            tipo_key = (conf.tipo or "outros").lower()
            if tipo_key in {"lote", "matricula"}:
                lotes_normais.append(conf)
            else:
                especiais.append(conf)
        lotes_normais.sort(key=lambda c: ((c.direcao or "").lower(), c.identificador or ""))
        especiais.sort(key=lambda c: ((c.tipo or "").lower(), c.identificador or ""))

        lotes_inseridos = 0
        for conf in lotes_normais:
            owners = owners_for(conf.matricula_anexada)
            owners_text = join_with_overflow(owners)
            matricula_value = conf.matricula_anexada or "‚ö†Ô∏è Matr√≠cula n√£o anexada"
            if not conf.matricula_anexada:
                ident = (conf.identificador or "Confrontante sem identifica√ß√£o").strip()
                if ident and ident not in missing_confrontantes:
                    missing_confrontantes.append(ident)
                if not owners_text:
                    owners_text = "‚ö†Ô∏è Propriet√°rios n√£o identificados (matr√≠cula n√£o anexada)"
            self.tree_confrontantes_lotes.insert(
                "",
                "end",
                values=(
                    format_direction(conf.direcao),
                    conf.identificador or "‚Äî",
                    matricula_value,
                    owners_text or "",
                ),
            )
            lotes_inseridos += 1

        if lotes_inseridos == 0 and matriculas_confrontantes_lista:
            for mat_num in matriculas_confrontantes_lista:
                mat_obj = matriculas_map.get(mat_num)
                if mat_obj and mat_obj.lote:
                    identificador = f"Lote {mat_obj.lote}"
                    if mat_obj.quadra:
                        identificador += f" / Quadra {mat_obj.quadra}"
                else:
                    identificador = f"Matr√≠cula {mat_num}"
                owners = owners_for(mat_num)
                owners_text = join_with_overflow(owners)
                self.tree_confrontantes_lotes.insert(
                    "",
                    "end",
                    values=("‚Äî", identificador, mat_num, owners_text),
                )
                lotes_inseridos += 1

        if lotes_inseridos == 0:
            self._insert_placeholder_row(self.tree_confrontantes_lotes, "Sem confrontantes de lote vinculados.")

        tipo_labels = {
            "via_publica": "Via p√∫blica",
            "via": "Via p√∫blica",
            "estado": "Estado",
            "pessoa": "Pessoa",
            "curso_dagua": "Curso d'√°gua",
            "rio": "Curso d'√°gua",
            "outros": "Outros",
        }

        especiais_inseridos = 0
        for conf in especiais:
            tipo_key = (conf.tipo or "outros").lower()
            tipo_label = tipo_labels.get(tipo_key, tipo_key.replace("_", " ").title())
            detalhes_parts = []
            if conf.matricula_anexada:
                detalhes_parts.append(f"Matr√≠cula anexada: {conf.matricula_anexada}")
                owners = owners_for(conf.matricula_anexada)
                owners_text = join_with_overflow(owners, limit=2)
                if owners_text:
                    detalhes_parts.append(f"Propriet√°rios: {owners_text}")
            elif tipo_key == "estado" and result.resumo_analise and result.resumo_analise.estado_ms_direitos:
                if result.resumo_analise.estado_ms_direitos.tem_direitos:
                    detalhes_parts.append("Direitos registrados em favor do Estado")
                elif result.resumo_analise.estado_ms_direitos.detalhes:
                    detalhes_parts.append("Direitos estadual sem confronta√ß√£o direta")
            detalhes = " | ".join(detalhes_parts)

            self.tree_confrontantes_outros.insert(
                "",
                "end",
                values=(
                    tipo_label,
                    conf.identificador or "‚Äî",
                    format_direction(conf.direcao),
                    detalhes,
                ),
            )
            especiais_inseridos += 1

        if especiais_inseridos == 0:
            self._insert_placeholder_row(self.tree_confrontantes_outros, "Sem confrontantes especiais identificados.")

        nao_conf_inseridos = 0
        for mat_num in matriculas_nao_confrontantes_lista:
            mat_obj = matriculas_map.get(mat_num)
            if mat_obj and mat_obj.lote:
                identificador = f"Lote {mat_obj.lote}"
                if mat_obj.quadra:
                    identificador += f" / Quadra {mat_obj.quadra}"
            else:
                identificador = f"Matr√≠cula {mat_num}"
            owners = owners_for(mat_num)
            owners_text = join_with_overflow(owners)
            self.tree_nao_confrontantes.insert(
                "",
                "end",
                values=(mat_num, identificador, owners_text),
            )
            nao_conf_inseridos += 1

        if nao_conf_inseridos == 0:
            self._insert_placeholder_row(self.tree_nao_confrontantes, "Sem matr√≠culas n√£o confrontantes anexadas.")

        if hasattr(self, "results_notebook"):
            self.results_notebook.select(self.tab_principal)

    def update_summary(self, result):
        """Atualiza o campo de resumo com o reasoning do modelo"""
        if not result:
            self.set_summary_text("Nenhuma an√°lise dispon√≠vel.")
            return
        
        # Usa o reasoning do modelo se dispon√≠vel
        if result.reasoning and result.reasoning.strip():
            # Adiciona informa√ß√µes b√°sicas + reasoning do modelo
            confianca = int(result.confidence * 100) if result.confidence is not None and result.confidence <= 1 else int(result.confidence) if result.confidence is not None else 0
            
            resumo_header = f"AN√ÅLISE PERICIAL (Confian√ßa: {confianca}%)\n\n"
            reasoning_texto = result.reasoning.strip()
            
            # Formata o reasoning para melhor legibilidade
            if reasoning_texto and not reasoning_texto.startswith("üìã"):
                reasoning_texto = f"üìã {reasoning_texto}"
            
            resumo = resumo_header + reasoning_texto
            self.set_summary_text(resumo)
        else:
            # Fallback para resumo autom√°tico se n√£o houver reasoning
            self._generate_fallback_summary(result)
    
    def _generate_fallback_summary(self, result):
        """Gera resumo autom√°tico caso n√£o haja reasoning do modelo"""
        if not result.matricula_principal:
            self.set_summary_text("Dados insuficientes para an√°lise.")
            return

        # Encontra dados da matr√≠cula principal
        matricula_principal_obj = None
        for mat in result.matriculas_encontradas:
            if mat.numero == result.matricula_principal:
                matricula_principal_obj = mat
                break

        if not matricula_principal_obj:
            self.set_summary_text("Dados da matr√≠cula principal n√£o encontrados.")
            return

        # Dados b√°sicos
        proprietarios = " e ".join(matricula_principal_obj.proprietarios) if matricula_principal_obj.proprietarios else "N√£o identificado"
        confianca = int(result.confidence * 100) if result.confidence is not None and result.confidence <= 1 else int(result.confidence) if result.confidence is not None else 0

        # Informa√ß√µes do lote/quadra
        lote_quadra = ""
        if matricula_principal_obj.lote or matricula_principal_obj.quadra:
            lote_parts = []
            if matricula_principal_obj.lote:
                lote_parts.append(f"Lote {matricula_principal_obj.lote}")
            if matricula_principal_obj.quadra:
                lote_parts.append(f"Quadra {matricula_principal_obj.quadra}")
            lote_quadra = f" ({', '.join(lote_parts)})"

        # An√°lise de confrontantes
        total_confrontantes = len(result.lotes_confrontantes) if result.lotes_confrontantes else 0
        confrontacao_adequada = "‚úÖ Adequada" if total_confrontantes >= 4 else f"‚ö†Ô∏è Insuficiente ({total_confrontantes} de 4 m√≠nimos)"

        # Confrontantes por tipo
        tipos_confrontantes = {}
        if result.lotes_confrontantes:
            for conf in result.lotes_confrontantes:
                tipo = conf.tipo
                if tipo not in tipos_confrontantes:
                    tipos_confrontantes[tipo] = 0
                tipos_confrontantes[tipo] += 1

        # Estado de MS
        estado_ms_confrontante = False
        estado_ms_direitos = False

        if result.lotes_confrontantes:
            for conf in result.lotes_confrontantes:
                if 'estado' in conf.identificador.lower() or 'mato grosso' in conf.identificador.lower():
                    estado_ms_confrontante = True
                    break

        if result.resumo_analise and hasattr(result.resumo_analise, 'estado_ms_direitos'):
            estado_ms_direitos = result.resumo_analise.estado_ms_direitos.get('tem_direitos', False)

        # Cadeia dominial
        cadeia_info = ""
        if result.resumo_analise and hasattr(result.resumo_analise, 'cadeia_dominial_completa'):
            cadeia_data = result.resumo_analise.cadeia_dominial_completa
            if cadeia_data and result.matricula_principal in cadeia_data:
                historico = cadeia_data[result.matricula_principal]
                if historico and len(historico) > 1:
                    cadeia_info = f"\nüìã Cadeia dominial: {len(historico)} transmiss√µes identificadas"

        # Restri√ß√µes
        restricoes_info = ""
        if result.resumo_analise:
            restricoes_vigentes = getattr(result.resumo_analise, 'restricoes_vigentes', [])
            restricoes_baixadas = getattr(result.resumo_analise, 'restricoes_baixadas', [])

            if restricoes_vigentes:
                restricoes_info += f"\n‚ö†Ô∏è {len(restricoes_vigentes)} restri√ß√£o(√µes) vigente(s)"
            if restricoes_baixadas:
                restricoes_info += f"\n‚úÖ {len(restricoes_baixadas)} restri√ß√£o(√µes) baixada(s)"

        # Monta o resumo completo
        resumo_partes = [
            f"üéØ RESUMO DA AN√ÅLISE (Confian√ßa: {confianca}%)",
            "",
            f"üìã MATR√çCULA PRINCIPAL: {result.matricula_principal}{lote_quadra}",
            f"üë§ PROPRIET√ÅRIO(S): {proprietarios}",
            "",
            f"üß≠ CONFRONTA√á√ÉO: {confrontacao_adequada}",
            f"üìä Total de confrontantes: {total_confrontantes}"
        ]

        # Adiciona detalhes dos tipos de confrontantes
        if tipos_confrontantes:
            resumo_partes.append("üìç Tipos identificados:")
            for tipo, qtd in tipos_confrontantes.items():
                emoji = {'lote': 'üèòÔ∏è', 'via_publica': 'üõ£Ô∏è', 'estado': 'üèõÔ∏è', 'pessoa': 'üë§', 'outros': 'üìç'}.get(tipo, 'üìç')
                resumo_partes.append(f"   {emoji} {tipo.replace('_', ' ').title()}: {qtd}")

        # Informa√ß√µes sobre propriet√°rios dos lotes confrontantes
        lotes_com_proprietarios = []
        if result.lotes_confrontantes:
            for conf in result.lotes_confrontantes:
                if conf.tipo in ['lote', 'matricula'] and conf.matricula_anexada:
                    # Encontra dados da matr√≠cula confrontante
                    confrontante_obj = None
                    for mat in result.matriculas_encontradas:
                        if mat.numero == conf.matricula_anexada:
                            confrontante_obj = mat
                            break

                    if confrontante_obj and confrontante_obj.proprietarios and confrontante_obj.proprietarios[0] != "N/A":
                        proprietarios_texto = ", ".join(confrontante_obj.proprietarios) if len(confrontante_obj.proprietarios) <= 2 else f"{confrontante_obj.proprietarios[0]} e mais {len(confrontante_obj.proprietarios)-1}"
                        lotes_com_proprietarios.append(f"   ‚Ä¢ {conf.identificador}: {proprietarios_texto}")

        if lotes_com_proprietarios:
            resumo_partes.append("")
            resumo_partes.append("üë• PROPRIET√ÅRIOS DOS LOTES CONFRONTANTES:")
            resumo_partes.extend(lotes_com_proprietarios[:5])  # M√°ximo 5 para n√£o poluir
            if len(lotes_com_proprietarios) > 5:
                resumo_partes.append(f"   ‚Ä¢ ... e mais {len(lotes_com_proprietarios) - 5} lotes")

        # Status do Estado de MS
        resumo_partes.append("")
        if estado_ms_confrontante:
            resumo_partes.append("üèõÔ∏è Estado de MS: ‚úÖ Identificado como confrontante")
        elif estado_ms_direitos:
            resumo_partes.append("üèõÔ∏è Estado de MS: ‚ö†Ô∏è Possui direitos registrados")
        else:
            resumo_partes.append("üèõÔ∏è Estado de MS: ‚úÖ N√£o identificado")

        # Adiciona informa√ß√µes complementares
        if cadeia_info:
            resumo_partes.append(cadeia_info)

        if restricoes_info:
            resumo_partes.append(restricoes_info)

        # Lotes n√£o confrontantes
        if result.matriculas_nao_confrontantes:
            lotes_nao_confrontantes_info = []
            for mat_num in result.matriculas_nao_confrontantes:
                # Encontra dados da matr√≠cula n√£o confrontante
                nao_confrontante_obj = None
                for mat in result.matriculas_encontradas:
                    if mat.numero == mat_num:
                        nao_confrontante_obj = mat
                        break

                if nao_confrontante_obj and nao_confrontante_obj.proprietarios and nao_confrontante_obj.proprietarios[0] != "N/A":
                    identificador = f"Lote {nao_confrontante_obj.lote}" if nao_confrontante_obj.lote else f"Matr√≠cula {mat_num}"
                    proprietarios_texto = ", ".join(nao_confrontante_obj.proprietarios) if len(nao_confrontante_obj.proprietarios) <= 2 else f"{nao_confrontante_obj.proprietarios[0]} e mais {len(nao_confrontante_obj.proprietarios)-1}"
                    lotes_nao_confrontantes_info.append(f"   ‚Ä¢ {identificador}: {proprietarios_texto}")

            if lotes_nao_confrontantes_info:
                resumo_partes.append("")
                resumo_partes.append("üìã LOTES N√ÉO CONFRONTANTES ANEXADOS:")
                resumo_partes.extend(lotes_nao_confrontantes_info[:3])  # M√°ximo 3
                if len(lotes_nao_confrontantes_info) > 3:
                    resumo_partes.append(f"   ‚Ä¢ ... e mais {len(lotes_nao_confrontantes_info) - 3} lotes")

        # Matr√≠culas encontradas
        resumo_partes.append(f"\nüìÑ Total de matr√≠culas analisadas: {len(result.matriculas_encontradas)}")

        resumo = "\n".join(resumo_partes)
        self.set_summary_text(resumo)

    def set_summary_text(self, text):
        """Atualiza o texto do campo de resumo"""
        self.txt_resumo.config(state=tk.NORMAL)
        self.txt_resumo.delete("1.0", tk.END)
        self.txt_resumo.insert("1.0", text)
        self.txt_resumo.config(state=tk.DISABLED)

    def diagnose_file_issues(self, file_path: str) -> str:
        """Diagnostica problemas comuns com arquivos"""
        issues = []
        
        try:
            # Verifica se arquivo existe
            if not os.path.exists(file_path):
                issues.append("‚ùå Arquivo n√£o encontrado")
                return "; ".join(issues)
            
            # Verifica tamanho do arquivo (apenas informativo)
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > 100:  # Aumentou limite
                issues.append(f"‚ÑπÔ∏è Arquivo grande ({file_size_mb:.1f}MB) - processamento pode demorar")
            
            # Verifica extens√£o
            ext = os.path.splitext(file_path.lower())[1]
            if ext == ".pdf":
                # Verifica n√∫mero de p√°ginas
                try:
                    page_count = get_pdf_page_count(file_path)
                    if page_count > 200:  # Limite muito alto, apenas informativo
                        issues.append(f"‚ÑπÔ∏è PDF com {page_count} p√°ginas - processamento pode demorar")
                    elif page_count == 0:
                        issues.append("‚ùå PDF corrompido ou vazio")
                except:
                    issues.append("‚ùå Erro ao ler PDF")
            elif ext not in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"]:
                issues.append(f"‚ùå Formato n√£o suportado: {ext}")
            
            if not issues:
                issues.append("‚úÖ Arquivo parece estar OK")
                
        except Exception as e:
            issues.append(f"‚ùå Erro ao analisar arquivo: {str(e)[:50]}")
        
        return "; ".join(issues)

    def check_estado_ms_rights(self, analysis_result: AnalysisResult) -> Optional[str]:
        """Verifica se o Estado de MS tem direitos registrados nas matr√≠culas"""
        direitos_encontrados = []
        
        # Verifica em todas as matr√≠culas
        for matricula in analysis_result.matriculas_encontradas:
            # Verifica se Estado de MS √© propriet√°rio
            for proprietario in matricula.proprietarios:
                if any(palavra in proprietario.lower() for palavra in 
                      ['estado de mato grosso do sul', 'estado de ms', 'estado do ms', 
                       'fazenda p√∫blica', 'governo do estado']):
                    direitos_encontrados.append(f"Matr√≠cula {matricula.numero}: Propriet√°rio")
            
            # Verifica restri√ß√µes onde Estado de MS √© credor
            for restricao in matricula.restricoes:
                if restricao.credor and any(palavra in restricao.credor.lower() for palavra in 
                                          ['estado de mato grosso do sul', 'estado de ms', 'estado do ms', 
                                           'fazenda p√∫blica', 'governo do estado']):
                    direitos_encontrados.append(
                        f"Matr√≠cula {matricula.numero}: {restricao.tipo.upper()} "
                        f"({restricao.situacao})"
                    )
        
        # Verifica resumo da an√°lise
        if analysis_result.resumo_analise:
            # Verifica estrutura espec√≠fica de direitos do Estado de MS
            if analysis_result.resumo_analise.estado_ms_direitos.tem_direitos:
                for detalhe in analysis_result.resumo_analise.estado_ms_direitos.detalhes:
                    direitos_encontrados.append(
                        f"‚ö†Ô∏è {detalhe.get('tipo_direito', 'Direito').upper()} "
                        f"(Status: {detalhe.get('status', 'N/A')})"
                    )
            
            # Verifica tamb√©m nas restri√ß√µes gerais
            for restricao in analysis_result.resumo_analise.restricoes_vigentes:
                if restricao.get('credor') and any(palavra in restricao['credor'].lower() for palavra in 
                                                 ['estado de mato grosso do sul', 'estado de ms', 'estado do ms', 
                                                  'fazenda p√∫blica', 'governo do estado']):
                    direitos_encontrados.append(
                        f"VIGENTE: {restricao.get('tipo', 'Restri√ß√£o').upper()}"
                    )
        
        if direitos_encontrados:
            return " | ".join(direitos_encontrados)
        return None

    def update_estado_alert(self):
        """Atualiza o alerta sobre direitos do Estado de MS"""
        direitos_estado = []
        
        # Verifica todos os resultados analisados
        for file_path, result in self.results.items():
            direitos = self.check_estado_ms_rights(result)
            if direitos:
                filename = os.path.basename(file_path)
                direitos_estado.append(f"{filename}: {direitos}")
        
        if direitos_estado:
            alert_text = "ATEN√á√ÉO: Estado de MS tem direitos registrados!\n" + "\n".join(direitos_estado)
            self.estado_alert_var.set(alert_text)
            self.estado_alert_label.pack(fill="x", pady=(0,5))
            # Piscar o alerta para chamar aten√ß√£o
            self.blink_alert()
        else:
            self.estado_alert_label.pack_forget()

    def blink_alert(self):
        """Faz o alerta piscar para chamar aten√ß√£o"""
        current_bg = self.estado_alert_label.cget("background")
        if current_bg == "yellow":
            self.estado_alert_label.configure(background="red", foreground="white")
            self.after(500, lambda: self.estado_alert_label.configure(background="yellow", foreground="red"))
        
        # Repete o piscar 3 vezes
        self.after(1000, self.blink_alert_cycle)

    def blink_alert_cycle(self):
        """Controla o ciclo de piscar do alerta"""
        if not hasattr(self, '_blink_count'):
            self._blink_count = 0
        
        if self._blink_count < 3:
            self.blink_alert()
            self._blink_count += 1
        else:
            self._blink_count = 0

    def generate_full_report(self):
        """Solicita √† LLM um relat√≥rio completo com todos os dados extra√≠dos."""
        if self.cached_full_report_text and self.cached_full_report_payload:
            self.log("üìÑ Reabrindo relat√≥rio completo j√° gerado.")
            self._show_full_report_window(self.cached_full_report_text, self.cached_full_report_payload, None)
            return

        if not self.results:
            messagebox.showwarning("Nenhum resultado", "Processe pelo menos um arquivo antes de gerar o relat√≥rio completo.")
            return

        model = FULL_REPORT_MODEL

        result = list(self.results.values())[-1]

        self.btn_full_report.config(state="disabled")
        self.log("üìù Gerando relat√≥rio completo com IA...")

        progress_window = tk.Toplevel(self)
        progress_window.title("Gerando Relat√≥rio Completo")
        progress_window.geometry("420x160")
        progress_window.transient(self)
        progress_window.grab_set()

        ttk.Label(progress_window, text="üßæ Elaborando relat√≥rio completo da matr√≠cula...").pack(pady=20)
        progress_bar = ttk.Progressbar(progress_window, mode="indeterminate")
        progress_bar.pack(pady=10, padx=20, fill="x")
        progress_bar.start()

        def run_generation():
            try:
                payload_dict = self._build_full_report_payload(result, model)
                payload_json = json.dumps(payload_dict, ensure_ascii=False, indent=2)
                prompt = build_full_report_prompt(payload_json)
                report_text = self._request_full_report(model, prompt)
                self.cached_full_report_text = report_text.strip()
                self.cached_full_report_payload = payload_json
                self.after(0, lambda: self._show_full_report_window(report_text, payload_json, progress_window))
            except Exception as exc:
                self.after(0, lambda: self._show_full_report_error(str(exc), progress_window))
            finally:
                self.after(0, lambda: self.btn_full_report.config(state="normal"))

        threading.Thread(target=run_generation, daemon=True).start()

    def _request_full_report(self, model: str, prompt: str) -> str:
        """Encapsula a chamada textual √† OpenRouter adicionando tratamento de erro contextual."""
        try:
            return call_openrouter_text(
                model=model,
                system_prompt=SYSTEM_PROMPT,
                user_prompt=prompt,
                temperature=0.2,
                max_tokens=3200
            )
        except Exception as exc:
            raise RuntimeError(f"Erro ao solicitar relat√≥rio completo: {exc}") from exc

    def _build_full_report_payload(self, result: AnalysisResult, model: str) -> Dict:
        """Agrupa todos os dados dispon√≠veis para envio ao modelo."""
        payload: Dict[str, Union[str, float, bool, List, Dict]] = {}
        if isinstance(result.raw_json, dict) and result.raw_json:
            payload["raw_json"] = result.raw_json

        payload.update({
            "arquivo": result.arquivo,
            "modelo_utilizado": model,
            "gerado_em": datetime.now().isoformat(),
            "matricula_principal": result.matricula_principal,
            "matriculas_confrontantes": result.matriculas_confrontantes,
            "matriculas_nao_confrontantes": result.matriculas_nao_confrontantes,
            "lotes_sem_matricula": result.lotes_sem_matricula,
            "confrontacao_completa": result.confrontacao_completa,
            "proprietarios_identificados": result.proprietarios_identificados,
            "confidence": result.confidence,
            "reasoning": result.reasoning,
            "estado_ms_inferido": result.is_confrontante,
        })

        payload["matriculas_encontradas"] = [asdict(mat) for mat in result.matriculas_encontradas]
        payload["lotes_confrontantes"] = [asdict(conf) for conf in result.lotes_confrontantes]
        payload["resumo_analise"] = asdict(result.resumo_analise) if result.resumo_analise else {}

        resumo_textual = self.txt_resumo.get("1.0", tk.END).strip()
        if resumo_textual:
            payload["resumo_textual"] = resumo_textual

        referencia_usuario = self.matricula_var.get().strip()
        if referencia_usuario and referencia_usuario != "ex: 12345":
            ref_value = referencia_usuario
        else:
            ref_value = ""
        payload["entrada_usuario"] = {
            "matricula_referencia_informada": ref_value,
            "arquivos_processados": [os.path.basename(path) for path in self.results.keys()]
        }

        missing_sem_matricula: List[str] = []
        for conf in result.lotes_confrontantes or []:
            tipo_key = (conf.tipo or "").lower()
            if tipo_key in {"lote", "matricula"} and not conf.matricula_anexada:
                ident = (conf.identificador or "").strip()
                if ident and ident not in missing_sem_matricula:
                    missing_sem_matricula.append(ident)
        for ident in result.lotes_sem_matricula or []:
            ident = (ident or "").strip()
            if ident and ident not in missing_sem_matricula:
                missing_sem_matricula.append(ident)
        if missing_sem_matricula:
            payload["confrontantes_sem_matricula"] = missing_sem_matricula

        payload["metricas"] = {
            "total_matriculas": len(result.matriculas_encontradas or []),
            "total_confrontantes_relacionados": len(result.lotes_confrontantes or []) or len(result.matriculas_confrontantes or []),
            "total_nao_confrontantes": len(result.matriculas_nao_confrontantes or []),
        }

        return payload

    def _show_full_report_window(self, report_text: str, payload_json: str, progress_window: Optional[tk.Toplevel]):
        """Exibe o relat√≥rio completo retornado pela IA."""
        from_cache = progress_window is None or not progress_window.winfo_exists()
        if progress_window and progress_window.winfo_exists():
            progress_window.destroy()

        if not from_cache:
            self.log("‚úÖ Relat√≥rio completo gerado.")

        report_window = tk.Toplevel(self)
        report_window.title("Relat√≥rio Completo da Matr√≠cula")
        report_window.geometry("920x740")
        report_window.transient(self)

        ttk.Label(report_window, text="üìÑ Relat√≥rio completo gerado pela IA", font=("Arial", 14, "bold")).pack(pady=(12, 6))

        text_frame = ttk.Frame(report_window)
        text_frame.pack(fill="both", expand=True, padx=12, pady=12)

        txt_report = tk.Text(
            text_frame,
            wrap="word",
            bg="white",
            relief="flat",
            padx=16,
            pady=12
        )
        txt_report.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=txt_report.yview)
        scrollbar.pack(side="right", fill="y")
        txt_report.configure(yscrollcommand=scrollbar.set)

        content = report_text.strip() or "Relat√≥rio vazio retornado pela IA."
        self._render_markdown_content(txt_report, content)

        button_frame = ttk.Frame(report_window)
        button_frame.pack(fill="x", padx=12, pady=(0, 12))

        # Primeira linha de bot√µes
        top_buttons = ttk.Frame(button_frame)
        top_buttons.pack(fill="x", pady=(0, 4))

        ttk.Button(top_buttons, text="üìã Copiar formatado", command=lambda: self._copy_formatted_to_clipboard(content)).pack(side="left")
        ttk.Button(top_buttons, text="üíæ Salvar", command=lambda: self._save_report_with_format_selection(content)).pack(side="left", padx=(8, 0))

        # Segunda linha de bot√µes
        bottom_buttons = ttk.Frame(button_frame)
        bottom_buttons.pack(fill="x")

        ttk.Button(bottom_buttons, text="üßæ Ver dados enviados", command=lambda: self._show_payload_window("Dados estruturados enviados", payload_json)).pack(side="left")
        ttk.Button(bottom_buttons, text="üíæ Salvar JSON", command=lambda: self._save_text_to_file(payload_json, title="Salvar Dados Estruturados", default_extension=".json")).pack(side="left", padx=(8, 0))
        ttk.Button(bottom_buttons, text="Fechar", command=report_window.destroy).pack(side="right")

    def _show_full_report_error(self, error: str, progress_window: tk.Toplevel):
        """Trata falhas na gera√ß√£o do relat√≥rio completo."""
        if progress_window and progress_window.winfo_exists():
            progress_window.destroy()
        self.log(f"‚ùå Erro ao gerar relat√≥rio completo: {error}")
        mensagem = f"N√£o foi poss√≠vel gerar o relat√≥rio completo:\n{error}"
        messagebox.showerror("Erro na gera√ß√£o", mensagem)
        self.cached_full_report_text = None
        self.cached_full_report_payload = None

    def _show_payload_window(self, title: str, content: str):
        """Abre janela modal exibindo texto estruturado (JSON)."""
        payload_window = tk.Toplevel(self)
        payload_window.title(title)
        payload_window.geometry("760x520")
        payload_window.transient(self)

        text_widget = tk.Text(payload_window, wrap="word", font=("Consolas", 10))
        text_widget.pack(fill="both", expand=True, padx=10, pady=10)
        text_widget.insert("1.0", content)
        text_widget.configure(state="disabled")

        button_frame = ttk.Frame(payload_window)
        button_frame.pack(fill="x", padx=10, pady=(0, 10))
        ttk.Button(button_frame, text="üìã Copiar", command=lambda: self._copy_to_clipboard(content)).pack(side="left")
        ttk.Button(button_frame, text="Fechar", command=payload_window.destroy).pack(side="right")

    def _setup_report_text_tags(self, widget: tk.Text):
        """Configura estilos para renderiza√ß√£o rica do relat√≥rio."""
        if getattr(widget, "_report_tags_configured", False):
            return

        try:
            tkfont.Font(family="Calibri", size=12)
            family = "Calibri"
            size = 11
        except tk.TclError:
            base_font = tkfont.nametofont("TkDefaultFont")
            family = base_font.actual("family")
            size = base_font.actual("size")

        widget.configure(font=(family, size), spacing3=8, bg="#fafafa")
        widget.tag_configure("paragraph", spacing3=12)
        widget.tag_configure("heading1", font=(family, size + 8, "bold"), spacing1=20, spacing3=16, foreground="#1a472a")
        widget.tag_configure("heading2", font=(family, size + 4, "bold"), spacing1=16, spacing3=12, foreground="#2c5530")
        widget.tag_configure("heading3", font=(family, size + 2, "bold"), spacing1=12, spacing3=10, foreground="#4a6741")
        widget.tag_configure("bold", font=(family, size, "bold"), foreground="#2c3e50")
        widget.tag_configure("bullet", lmargin1=40, lmargin2=60, spacing3=8, foreground="#34495e")
        widget.tag_configure("numbered", lmargin1=40, lmargin2=60, spacing3=8, foreground="#34495e")
        widget.tag_configure("hr", spacing3=12)

        widget._report_tags_configured = True

    def _insert_formatted_text(self, widget: tk.Text, text: str, base_tags: Tuple[str, ...]):
        """Insere texto aplicando formata√ß√£o simples de negrito."""
        remainder = text
        while True:
            start = remainder.find("**")
            if start == -1:
                if remainder:
                    widget.insert(tk.END, remainder, base_tags)
                break
            if start > 0:
                widget.insert(tk.END, remainder[:start], base_tags)
            remainder = remainder[start + 2:]
            end = remainder.find("**")
            if end == -1:
                widget.insert(tk.END, "**" + remainder, base_tags)
                break
            bold_segment = remainder[:end]
            widget.insert(tk.END, bold_segment, base_tags + ("bold",))
            remainder = remainder[end + 2:]

    def _render_markdown_content(self, widget: tk.Text, markdown_text: str):
        """Renderiza conte√∫do markdown b√°sico com estilo de relat√≥rio."""
        self._setup_report_text_tags(widget)
        widget.configure(state=tk.NORMAL)
        widget.delete("1.0", tk.END)

        lines = markdown_text.splitlines()
        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()

            if not stripped:
                widget.insert(tk.END, "\n", ("paragraph",))
                continue

            if stripped.startswith("### "):
                heading = stripped[4:].strip()
                self._insert_formatted_text(widget, heading + "\n", ("heading3",))
                continue

            if stripped.startswith("## "):
                heading = stripped[3:].strip()
                self._insert_formatted_text(widget, heading.upper() + "\n", ("heading2",))
                continue

            if stripped.startswith("# "):
                heading = stripped[2:].strip()
                self._insert_formatted_text(widget, heading.upper() + "\n", ("heading1",))
                continue

            if stripped in {"---", "***"}:
                widget.insert(tk.END, "\n", ("hr",))
                continue

            bullet_prefixes = ("- ", "* ", "‚Ä¢ ")
            if stripped.startswith(bullet_prefixes):
                content = stripped[2:].strip()
                display = f"‚Ä¢ {content}\n"
                self._insert_formatted_text(widget, display, ("bullet",))
                continue

            number_split = stripped.split(". ", 1)
            if len(number_split) == 2 and number_split[0].isdigit():
                number, content = number_split
                display = f"{number}. {content.strip()}\n"
                self._insert_formatted_text(widget, display, ("numbered",))
                continue

            self._insert_formatted_text(widget, stripped + "\n", ("paragraph",))

        widget.configure(state=tk.DISABLED)
        widget.see("1.0")

    def _copy_to_clipboard(self, text: str):
        """Copia texto para a √°rea de transfer√™ncia"""
        self.clipboard_clear()
        self.clipboard_append(text)
        self.update()
        messagebox.showinfo("Copiado", "Conte√∫do copiado para a √°rea de transfer√™ncia!")

    def _copy_formatted_to_clipboard(self, markdown_text: str):
        """Copia texto formatado (RTF) para a √°rea de transfer√™ncia para preservar formata√ß√£o no Word"""
        try:
            rtf_text = self._markdown_to_rtf(markdown_text)
            self.clipboard_clear()
            self.clipboard_append(rtf_text)
            self.update()
            messagebox.showinfo("Copiado", "Conte√∫do formatado copiado para a √°rea de transfer√™ncia!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao copiar formata√ß√£o: {e}")
            # Fallback para texto simples
            self._copy_to_clipboard(markdown_text)

    def _save_text_to_file(self, text: str, title: str = "Salvar arquivo de texto", default_extension: str = ".txt"):
        """Abre di√°logo para salvar conte√∫do em arquivo."""
        filename = filedialog.asksaveasfilename(
            defaultextension=default_extension,
            filetypes=[("Text files", "*.txt"), ("JSON", "*.json"), ("All files", "*.*")],
            title=title
        )
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(text)
                messagebox.showinfo("Salvo", f"Arquivo salvo em: {filename}")
            except Exception as exc:
                messagebox.showerror("Erro ao salvar", f"N√£o foi poss√≠vel salvar o arquivo: {exc}")

    def _markdown_to_rtf(self, markdown_text: str) -> str:
        """Converte markdown simples para RTF b√°sico"""
        rtf = r"{\rtf1\ansi\deff0 {\fonttbl\f0\froman\fcharset0 Times New Roman;}"

        lines = markdown_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                rtf += r"\par "
                continue

            if line.startswith('# '):
                title = line[2:].strip()
                title = self._process_bold_for_rtf(title)
                rtf += rf"\fs32\b {title}\b0\fs24\par "
            elif line.startswith('## '):
                subtitle = line[3:].strip()
                subtitle = self._process_bold_for_rtf(subtitle)
                rtf += rf"\fs28\b {subtitle}\b0\fs24\par "
            elif line.startswith('### '):
                heading = line[4:].strip()
                heading = self._process_bold_for_rtf(heading)
                rtf += rf"\fs26\b {heading}\b0\fs24\par "
            elif line.startswith('‚Ä¢ ') or line.startswith('- '):
                content = line[2:].strip()
                content = self._process_bold_for_rtf(content)
                rtf += rf"\li360 \bullet {content}\par "
            else:
                content = self._process_bold_for_rtf(line)
                rtf += rf"{content}\par "

        rtf += "}"
        return rtf

    def _process_bold_for_rtf(self, text: str) -> str:
        """Processa marca√ß√µes de negrito ** para RTF"""
        result = ""
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Texto normal
                result += part
            else:  # Texto em negrito
                result += rf"\b {part}\b0 "
        return result

    def _save_as_docx(self, markdown_text: str):
        """Salva o relat√≥rio como arquivo DOCX"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Erro", "Biblioteca python-docx n√£o instalada. Execute: pip install python-docx")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="Salvar Relat√≥rio como DOCX"
        )

        if not filename:
            return

        try:
            doc = Document()

            # Configurar estilos
            styles = doc.styles

            # Estilo para t√≠tulo principal
            if 'Titulo1' not in styles:
                title_style = styles.add_style('Titulo1', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.size = Pt(18)
                title_style.font.bold = True
                title_style.font.name = 'Calibri'

            # Estilo para subt√≠tulos
            if 'Titulo2' not in styles:
                subtitle_style = styles.add_style('Titulo2', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_style.font.size = Pt(14)
                subtitle_style.font.bold = True
                subtitle_style.font.name = 'Calibri'

            # Processa o markdown
            lines = markdown_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                if line.startswith('# '):
                    title = line[2:].strip()
                    doc.add_heading(title, level=1)
                elif line.startswith('## '):
                    subtitle = line[3:].strip()
                    doc.add_heading(subtitle, level=2)
                elif line.startswith('### '):
                    heading = line[4:].strip()
                    doc.add_heading(heading, level=3)
                elif line.startswith(('‚Ä¢ ', '- ', '* ')):
                    content = line[2:].strip()
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    self._add_formatted_text(p, content)
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, line)

            doc.save(filename)
            messagebox.showinfo("Salvo", f"Relat√≥rio DOCX salvo em: {filename}")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar DOCX: {e}")

    def _add_formatted_text(self, paragraph, text):
        """Adiciona texto com formata√ß√£o em negrito ao par√°grafo do Word"""
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Texto normal
                paragraph.add_run(part)
            else:  # Texto em negrito
                paragraph.add_run(part).bold = True

    def _save_as_pdf(self, markdown_text: str):
        """Salva o relat√≥rio como arquivo PDF"""
        if not REPORTLAB_AVAILABLE:
            messagebox.showerror("Erro", "Biblioteca reportlab n√£o instalada. Execute: pip install reportlab")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            title="Salvar Relat√≥rio como PDF"
        )

        if not filename:
            return

        try:
            doc = SimpleDocTemplate(filename, pagesize=A4)
            styles = getSampleStyleSheet()

            # Criar estilos customizados
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                alignment=TA_LEFT
            )

            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10,
                alignment=TA_LEFT
            )

            story = []
            lines = markdown_text.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue

                if line.startswith('# '):
                    title = line[2:].strip()
                    story.append(Paragraph(title, title_style))
                elif line.startswith('## '):
                    subtitle = line[3:].strip()
                    story.append(Paragraph(subtitle, subtitle_style))
                elif line.startswith('### '):
                    heading = line[4:].strip()
                    story.append(Paragraph(f"<b>{heading}</b>", styles['Heading3']))
                elif line.startswith(('‚Ä¢ ', '- ', '* ')):
                    content = line[2:].strip()
                    # Processa negrito
                    content = content.replace('**', '<b>').replace('**', '</b>')
                    story.append(Paragraph(f"‚Ä¢ {content}", styles['Normal']))
                else:
                    # Processa negrito
                    content = line.replace('**', '<b>').replace('**', '</b>')
                    story.append(Paragraph(content, styles['Normal']))

            doc.build(story)
            messagebox.showinfo("Salvo", f"Relat√≥rio PDF salvo em: {filename}")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar PDF: {e}")

    def _save_report_with_format_selection(self, markdown_text: str):
        """Permite ao usu√°rio escolher o formato para salvar o relat√≥rio (DOCX ou PDF)"""
        import tkinter.messagebox as messagebox
        from tkinter import simpledialog

        # Criar janela de sele√ß√£o de formato
        format_window = tk.Toplevel()
        format_window.title("Escolher Formato")
        format_window.geometry("300x150")
        format_window.resizable(False, False)
        format_window.transient()
        format_window.grab_set()

        # Centralizar janela
        format_window.update_idletasks()
        x = (format_window.winfo_screenwidth() // 2) - (300 // 2)
        y = (format_window.winfo_screenheight() // 2) - (150 // 2)
        format_window.geometry(f"300x150+{x}+{y}")

        selected_format = tk.StringVar()

        # T√≠tulo
        title_label = ttk.Label(format_window, text="Escolha o formato para salvar:", font=("Arial", 11, "bold"))
        title_label.pack(pady=(20, 15))

        # Frame para bot√µes
        button_frame = ttk.Frame(format_window)
        button_frame.pack(expand=True)

        def save_as_docx():
            selected_format.set("docx")
            format_window.destroy()

        def save_as_pdf():
            selected_format.set("pdf")
            format_window.destroy()

        # Bot√µes de formato
        ttk.Button(button_frame, text="üìù DOCX", command=save_as_docx, width=12).pack(side="left", padx=10)
        ttk.Button(button_frame, text="üìÑ PDF", command=save_as_pdf, width=12).pack(side="left", padx=10)

        # Bot√£o cancelar
        cancel_frame = ttk.Frame(format_window)
        cancel_frame.pack(side="bottom", pady=(10, 15))
        ttk.Button(cancel_frame, text="Cancelar", command=format_window.destroy).pack()

        # Aguardar sele√ß√£o
        format_window.wait_window()

        # Executar salvamento baseado na sele√ß√£o
        if selected_format.get() == "docx":
            self._save_as_docx(markdown_text)
        elif selected_format.get() == "pdf":
            self._save_as_pdf(markdown_text)

    def log(self, msg: str):
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")

# =========================
# Main
# =========================
def main():
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
